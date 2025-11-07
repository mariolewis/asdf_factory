# agents/agent_report_generator.py

"""
This module contains the ReportGeneratorAgent class.
This agent is responsible for generating formatted .docx reports.
"""

import logging
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from io import BytesIO
import json
import re
from asdf_db_manager import ASDFDBManager
from pathlib import Path
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from htmldocx import HtmlToDocx
from openpyxl import Workbook
from openpyxl.styles import Font
import pandas as pd
from datetime import datetime, timezone
from gui.utils import format_timestamp_for_display
from gui.rendering_utils import generate_mermaid_png, generate_plotly_png
import plotly.graph_objects as go
import plotly.io as pio
import markdown
from bs4 import BeautifulSoup

class ReportGeneratorAgent:
    """
    Agent responsible for generating downloadable .docx files for various
    project reports.
    """

    def __init__(self, db_manager: ASDFDBManager = None):
        """
        Initializes the ReportGeneratorAgent.

        Args:
            db_manager (ASDFDBManager, optional): An instance of the database manager.
                Required for template-based generation.
        """
        self.db_manager = db_manager
        # We check for the db_manager inside the methods to ensure
        # existing calls that don't need it (like for .xlsx) don't break.
        logging.info("ReportGeneratorAgent initialized.")

    def _add_styled_paragraph(self, document, text, style_name):
        """Helper to add a paragraph with style fallback."""
        if text.strip() == "": # Don't add empty paragraphs unless they are intentional
            return
        try:
            document.add_paragraph(text, style=style_name)
        except KeyError:
            logging.warning(f"Style '{style_name}' not found. Using 'Normal'.")
            try:
                document.add_paragraph(text, style='Normal')
            except KeyError:
                logging.warning("Style 'Normal' not found. Using default paragraph.")
                document.add_paragraph(text)

    def _get_styled_document(self) -> Document:
        """
        Helper method to load the user's selected .docx template or fall back
        to a default Document object.
        """
        template_path_str = "data/templates/styles/default_docx_template.docx" # Default

        if self.db_manager:
            try:
                config_path = self.db_manager.get_config_value("SELECTED_DOCX_STYLE_PATH")
                if config_path:
                    template_path_str = config_path
                else:
                    logging.warning("SELECTED_DOCX_STYLE_PATH was empty in DB. Reverting to default.")
                    template_path_str = "data/templates/styles/default_docx_template.docx"
            except Exception as e:
                logging.error(f"Failed to read SELECTED_DOCX_STYLE_PATH from DB: {e}. Reverting to default.")
                template_path_str = "data/templates/styles/default_docx_template.docx"
        else:
            logging.warning("db_manager is None. Cannot load custom docx template. Using default.")
            template_path_str = "data/templates/styles/default_docx_template.docx"

        template_path = Path(template_path_str)

        if template_path.exists():
            try:
                document = Document(template_path)
                logging.debug(f"Loaded .docx template: {template_path_str}")
                return document
            except Exception as e:
                logging.error(f"Failed to load .docx template '{template_path_str}': {e}. Using default.")
                return Document()
        else:
            logging.warning(f"Template file not found at '{template_path_str}'. Using default.")
            return Document()

    def generate_progress_summary_docx(self, total_components: int, status_counts: dict) -> BytesIO:
        """
        Generates a .docx file for the Development Progress Summary report,
        using the central styled template.
        """
        document = self._get_styled_document()

        try:
            document.add_paragraph('ASDF - Development Progress Summary', style='Title')
        except KeyError:
            logging.warning("Style 'Title' not found. Using fallback.")
            document.add_heading('ASDF - Development Progress Summary', level=1)

        try:
            document.add_paragraph(f"Total Components Defined in Plan: {total_components}", style='Normal')
            document.add_paragraph() # Spacer

            document.add_paragraph('Components by Status', style='Heading 2')
            if status_counts:
                for status, count in status_counts.items():
                    # Use the named style for bullet lists
                    document.add_paragraph(f"{status}: {count}", style='Bullet List')
            else:
                document.add_paragraph("No components with status available.", style='Normal')

        except KeyError as e:
            logging.warning(f"Style not found in template, using fallback. Error: {e}")
            # Fallback logic if styles are missing
            if 'Normal' in str(e):
                document.add_paragraph(f"Total Components Defined in Plan: {total_components}")
            if 'Heading 2' in str(e):
                document.add_heading('Components by Status', level=2)
            if 'Bullet List' in str(e) and status_counts:
                for status, count in status_counts.items():
                    document.add_paragraph(f"{status}: {count}", style='List Bullet') # Use docx built-in
            elif not status_counts:
                document.add_paragraph("No components with status available.")

        except Exception as e:
            logging.error(f"Error applying styles in generate_progress_summary_docx: {e}")
            document.add_paragraph(f"An unexpected error occurred: {e}")

        # Save document to an in-memory buffer
        doc_buffer = BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer

    def generate_cr_bug_report_docx(self, report_data: list, filter_type: str) -> BytesIO:
        """
        Generates a .docx file for the Change Requests & Bug Fixes report,
        using the central styled template.
        """
        document = self._get_styled_document()

        try:
            document.add_paragraph('ASDF - Change Requests & Bug Fixes', style='Title')
            document.add_paragraph(f"Report Filter: {filter_type}", style='Normal')
            document.add_paragraph() # Spacer
        except KeyError:
            logging.warning("Style not found. Using fallbacks.")
            document.add_heading('ASDF - Change Requests & Bug Fixes', level=1)
            document.add_paragraph(f"Report Filter: {filter_type}")

        try:
            if not report_data:
                document.add_paragraph("No items match the selected filter.", style='Normal')
            else:
                # This logic is unchanged. The table will use the template's 'Table Grid' style.
                headers = ["ID", "Type", "Status", "Description"]
                table = document.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, header_name in enumerate(headers):
                    hdr_cells[i].text = header_name

                # Populate table with data
                for item in report_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item.get("id", "N/A"))
                    row_cells[1].text = item.get("type", "N/A")
                    row_cells[2].text = item.get("status", "N/A")
                    row_cells[3].text = item.get("description", "N/A")

        except KeyError as e:
            logging.warning(f"Style 'Normal' not found, using fallback. Error: {e}")
            document.add_paragraph("No items match the selected filter.")
        except Exception as e:
            logging.error(f"Error applying styles/table in generate_cr_bug_report_docx: {e}")
            document.add_paragraph(f"An unexpected error occurred: {e}")

        # Save document to an in-memory buffer
        doc_buffer = BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer

    def generate_assessment_docx(self, analysis_data: dict, project_name: str) -> BytesIO:
        """
        Generates a formatted .docx file for the Complexity & Risk Assessment report
        by formatting Markdown and passing to the central styling method.
        """
        title = f"Complexity & Risk Assessment: {project_name}"
        md_content = []

        comp_analysis = analysis_data.get("complexity_analysis", {})
        if not comp_analysis:
             md_content.append("Could not parse the analysis result.")
        else:
            md_content.append("## Complexity Analysis")
            for key, value in comp_analysis.items():
                title_str = key.replace('_', ' ').title()
                rating = value.get('rating', 'N/A')
                justification = value.get('justification', 'No details provided.')
                md_content.append(f"### {title_str}: {rating}\n\n*{justification}*\n")

            md_content.append("## Risk Assessment")
            risk_assessment = analysis_data.get("risk_assessment", {})
            md_content.append(f"**Overall Risk Level:** {risk_assessment.get('overall_risk_level', 'N/A')}\n")
            md_content.append(f"**Summary:** {risk_assessment.get('summary', 'No summary provided.')}\n")
            md_content.append(f"**Token Consumption Outlook:** {risk_assessment.get('token_consumption_outlook', 'N/A')}\n")

            recommendations = risk_assessment.get('recommendations', [])
            if recommendations:
                md_content.append("\n### Recommendations\n")
                for rec in recommendations:
                    md_content.append(f"* {rec}")

        # Call the new central method
        return self.generate_text_document_docx(title, "\n".join(md_content), is_html=False)

    def generate_text_document_docx(self, title: str, content: str, is_html: bool = False) -> BytesIO:
        """
        Generates a generic .docx file for text-based project documents.
        This version converts Markdown to HTML, then parses the HTML
        to robustly apply named styles from a template.
        """
        document = self._get_styled_document()

        # 1. Add the main title
        try:
            document.add_paragraph(title, style='Title')
        except Exception:
            logging.warning("Style 'Title' not found. Using default Heading 1.")
            document.add_heading(title, level=1)

        document.add_paragraph()

        # 2. Check for pre-formatted HTML (for manual test plans)
        if is_html:
            try:
                parser = HtmlToDocx()
                parser.add_html_to_document(content, document)
            except Exception as e:
                logging.error(f"Failed to parse HTML for DOCX: {e}. Falling back to plain text.")
                document.add_paragraph(content, style='Normal')

        # 3. Standard Markdown Processing
        else:
            try:
                # 1. Convert Markdown to HTML with extensions for code blocks, tables, and sane lists
                html_content = markdown.markdown(content, extensions=['fenced_code', 'tables', 'sane_lists'])
                soup = BeautifulSoup(html_content, 'html.parser')

                root_tags = soup.find_all(True, recursive=False)
                search_node = soup
                if len(root_tags) == 1 and root_tags[0].name not in ['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'pre', 'hr']:
                    # If the root is a single, unknown tag (like <div>),
                    # search inside it instead of the document root.
                    search_node = root_tags[0]

                # 2. Iterate over all *top-level* tags and map them to Word styles
                for tag in search_node.find_all(True, recursive=False):
                    style_name = 'Normal' # Default

                    if tag.name == 'h1':
                        style_name = 'Heading 1'
                        self._add_styled_paragraph(document, tag.get_text(), style_name)

                    elif tag.name == 'h2':
                        style_name = 'Heading 2'
                        self._add_styled_paragraph(document, tag.get_text(), style_name)

                    elif tag.name == 'h3':
                        style_name = 'Heading 3'
                        self._add_styled_paragraph(document, tag.get_text(), style_name)

                    elif tag.name == 'p':
                        style_name = 'Normal'
                        self._add_styled_paragraph(document, tag.get_text(), style_name)

                    elif tag.name == 'ul':
                        style_name = 'Bullet List'
                        for li in tag.find_all('li'):
                            self._add_styled_paragraph(document, li.get_text(), style_name)

                    elif tag.name == 'ol':
                        style_name = 'List Numbered'
                        for li in tag.find_all('li'):
                            self._add_styled_paragraph(document, li.get_text(), style_name)

                    elif tag.name == 'pre':
                        # This handles both ```code``` and ```mermaid``` blocks
                        style_name = 'Code Block'
                        # .get_text() correctly extracts text from <code> child
                        self._add_styled_paragraph(document, tag.get_text(), style_name)

                    elif tag.name == 'hr':
                        document.add_page_break()

            except Exception as e:
                logging.error(f"Fatal error during Markdown-to-HTML parsing: {e}")
                # Fallback: just dump the raw text
                self._add_styled_paragraph(document, content, 'Normal')

        # 4. Save to in-memory buffer
        doc_buffer = BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer

    def generate_dev_plan_docx(self, plan_data: dict, project_name: str) -> BytesIO:
        """
        Generates a formatted .docx file for the Development Plan
        by formatting Markdown and passing to the central styling method.

        Args:
            plan_data (dict): The structured data for the development plan.
            project_name (str): The name of the project for the report title.

        Returns:
            BytesIO: An in-memory byte stream of the generated .docx file.
        """
        title = f"Development Plan: {project_name}"
        md_content = []

        main_exe = plan_data.get("main_executable_file", "Not specified")
        # This line is corrected to NOT have bolding, matching the original.
        md_content.append(f"Main Executable File: {main_exe}\n")

        md_content.append("## Development Steps")

        plan_steps = plan_data.get("development_plan", [])
        if not plan_steps:
            md_content.append("No development steps were generated.")
        else:
            for i, task in enumerate(plan_steps, 1):
                # This formatting matches the sample's pattern
                md_content.append(f"\n**Task {i}: {task.get('component_name', 'N/A')}**")
                md_content.append(f"Description: {task.get('task_description', 'No description.')}")
                md_content.append(f"ID: {task.get('micro_spec_id', 'N/A')}")

        # Combine the list into a single string
        content = "\n".join(md_content)

        # Call the new central method
        return self.generate_text_document_docx(title, content, is_html=False)

    def generate_sprint_summary_text(self, summary_data: dict) -> str:
        """
        Formats sprint summary data into a readable Markdown string for the UI.

        Args:
            summary_data (dict): The structured data from the orchestrator.

        Returns:
            A string containing the summary in Markdown format.
        """
        if "error" in summary_data:
            return f"### Error\nCould not generate sprint summary: {summary_data['error']}"

        sprint_goal = summary_data.get("sprint_goal", "Not defined.")
        completed_items = summary_data.get("completed_items", [])

        md_parts = [f"### Sprint Goal\n{sprint_goal}\n\n### Delivered Items"]

        if not completed_items:
            md_parts.append("\nNo items were marked as completed in this sprint.")
        else:
            for item in completed_items:
                md_parts.append(f"\n* **{item.get('hierarchical_id', 'ID?')} {item.get('title', 'No Title')}** - Status: {item.get('status', 'N/A')}")

        return "".join(md_parts)

    def generate_backlog_xlsx(self, backlog_data: list) -> BytesIO:
        """
        Generates an .xlsx file for the full project backlog.

        Args:
            backlog_data (list): The hierarchical list of backlog item dictionaries.

        Returns:
            BytesIO: An in-memory byte stream of the generated .xlsx file.
        """
        flat_list = []

        def flatten_hierarchy(items, prefix=""):
            for i, item in enumerate(items, 1):
                current_prefix = f"{prefix}{i}"

                timestamp_str = item.get('last_modified_timestamp') or item.get('creation_timestamp')
                formatted_date = format_timestamp_for_display(timestamp_str)

                record = {
                    '#': current_prefix,
                    'Title': item.get('title', 'N/A'),
                    'Type': item.get('request_type', 'N/A').replace('_', ' ').title(),
                    'Status': item.get('status', 'N/A'),
                    'Priority/Severity': item.get('priority') or item.get('impact_rating') or '',
                    'Complexity': item.get('complexity', ''),
                    'Last Modified': formatted_date
                }
                flat_list.append(record)

                if "features" in item:
                    flatten_hierarchy(item["features"], prefix=f"{current_prefix}.")
                if "user_stories" in item:
                    flatten_hierarchy(item["user_stories"], prefix=f"{current_prefix}.")

        flatten_hierarchy(backlog_data)

        df = pd.DataFrame(flat_list)

        # Save dataframe to an in-memory buffer
        output_buffer = BytesIO()
        with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Project Backlog')

        output_buffer.seek(0)
        return output_buffer

    def generate_sprint_plan_docx(self, project_name: str, sprint_items: list, plan_data: list) -> BytesIO:
        """
        Generates a formatted .docx file for the Sprint Plan
        by formatting Markdown and passing to the central styling method.
        """
        title = f"Sprint Plan: {project_name}"
        md_content = []

        timestamp_str = datetime.now(timezone.utc).isoformat()
        timestamp = format_timestamp_for_display(timestamp_str)
        md_content.append(f"Generated on: {timestamp}\n")

        md_content.append("## Sprint Scope")
        if not sprint_items:
            md_content.append("No items were included in this sprint.")
        else:
            for item in sprint_items:
                item_id = item.get('hierarchical_id', 'N/A')
                item_title = item.get('title', 'Untitled Item')
                md_content.append(f"* {item_id}: {item_title}")

        md_content.append("\n## Implementation Plan")
        if not plan_data or (isinstance(plan_data, list) and plan_data and plan_data[0].get("error")):
            md_content.append("No implementation plan was generated.")
        else:
            for i, task in enumerate(plan_data, 1):
                md_content.append(f"\n**Task {i}: {task.get('component_name', 'N/A')}**")
                md_content.append(f"File Path: {task.get('component_file_path', 'N/A')}")
                md_content.append(f"Description: {task.get('task_description', 'No description.')}")

        # Call the new central method
        return self.generate_text_document_docx(title, "\n".join(md_content), is_html=False)

    def generate_sprint_summary_docx(self, summary_data: dict, project_name: str) -> BytesIO:
        """
        Generates a formatted .docx file for the Sprint Review summary
        by formatting Markdown and passing to the central styling method.
        """
        title = f"Sprint Summary Report: {project_name}"
        md_content = []

        timestamp_str = datetime.now(timezone.utc).isoformat()
        timestamp = format_timestamp_for_display(timestamp_str)
        md_content.append(f"Generated on: {timestamp}\n")

        if "error" in summary_data:
            md_content.append(f"Could not generate sprint summary: {summary_data['error']}")
        else:
            sprint_goal = summary_data.get("sprint_goal", "Not defined.")
            completed_items = summary_data.get("completed_items", [])

            md_content.append("## Sprint Goal")
            md_content.append(sprint_goal)

            md_content.append("\n## Delivered Items")
            if not completed_items:
                md_content.append("No items were marked as completed in this sprint.")
            else:
                for item in completed_items:
                    md_content.append(f"* **{item.get('hierarchical_id', 'ID?')} {item.get('title', 'No Title')}** - Status: {item.get('status', 'N/A')}")

        # Call the new central method
        return self.generate_text_document_docx(title, "\n".join(md_content), is_html=False)

    def generate_pre_execution_report_docx(self, project_name: str, selected_items: list, report_data: dict) -> BytesIO:
        """
        Generates a formatted .docx file for the Pre-Execution Check report
        by formatting Markdown and passing to the central styling method.
        """
        title = f"Sprint Pre-Execution Check Report: {project_name}"
        md_content = []

        timestamp_str = datetime.now(timezone.utc).isoformat()
        timestamp = format_timestamp_for_display(timestamp_str)
        md_content.append(f"Generated on: {timestamp}\n")

        md_content.append("## Selected Items for Analysis")
        if not selected_items:
            md_content.append("No items were selected for analysis.")
        else:
            for item in selected_items:
                item_id = item.get('hierarchical_id', 'N/A')
                item_title = item.get('title', 'Untitled Item')
                md_content.append(f"* {item_id}: {item_title}")

        md_content.append("\n## AI Analysis Report")
        dependencies = report_data.get("missing_dependencies", [])
        conflicts = report_data.get("technical_conflicts", [])
        advice = report_data.get("sequencing_advice", [])

        if not dependencies and not conflicts and not advice:
            md_content.append("No potential issues were found.")

        if dependencies:
            md_content.append("\n### Missing Dependencies")
            for item in dependencies:
                md_content.append(f"* {item}")

        if conflicts:
            md_content.append("\n### Potential Technical Conflicts")
            for item in conflicts:
                md_content.append(f"* {item}")

        if advice:
            md_content.append("\n### Architectural Sequencing Advice")
            for item in advice:
                md_content.append(f"* {item}")

        # Call the new central method
        return self.generate_text_document_docx(title, "\n".join(md_content), is_html=False)

    def generate_traceability_docx(self, trace_data: list, project_name: str) -> BytesIO:
        """
        Generates a formatted .docx file for the Requirements Traceability report,
        using the central styled template.
        """
        logging.info(f"Generating DOCX for traceability report (styled) for project: {project_name}")
        document = self._get_styled_document()

        # Change orientation to Landscape (must be done *after* loading template)
        try:
            section = document.sections[0]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            # Set smaller margins for more table space
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
        except Exception as e:
            logging.warning(f"Could not set landscape orientation: {e}")

        # Add Title and Timestamp
        timestamp_str = datetime.now(timezone.utc).isoformat()
        timestamp = format_timestamp_for_display(timestamp_str)

        try:
            document.add_paragraph(f"Requirements Traceability Report: {project_name}", style='Title')
            document.add_paragraph(f"Generated on: {timestamp}", style='Normal')
        except KeyError:
            logging.warning("Style not found. Using fallbacks.")
            document.add_heading(f"Requirements Traceability Report: {project_name}", level=1)
            document.add_paragraph(f"Generated on: {timestamp}")

        document.add_paragraph() # Add a blank line

        try:
            if not trace_data:
                document.add_paragraph("No traceability data available for this project.", style='Normal')
            else:
                # Define table headers (matching the UI)
                headers = ['Backlog Item (#)', 'Title', 'Status', 'Artifact Path', 'Artifact Name']

                # Create Table (this logic is unchanged and will adopt template styles)
                table = document.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid' # Apply a basic style
                hdr_cells = table.rows[0].cells
                for i, header_name in enumerate(headers):
                    hdr_cells[i].text = header_name

                # Populate table with data
                for item in trace_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item.get("backlog_id", "N/A")
                    row_cells[1].text = item.get("backlog_title", "N/A")
                    row_cells[2].text = item.get("backlog_status", "N/A")
                    row_cells[3].text = item.get("artifact_path", "N/A")
                    row_cells[4].text = item.get("artifact_name", "N/A")

        except KeyError as e:
             logging.warning(f"Style 'Normal' not found, using fallback. Error: {e}")
             document.add_paragraph("No traceability data available for this project.")
        except Exception as table_e:
            logging.error(f"Error creating DOCX table for traceability: {table_e}", exc_info=True)
            document.add_paragraph(f"Error creating report table: {table_e}")

        # Save document to an in-memory buffer
        doc_buffer = BytesIO()
        try:
            document.save(doc_buffer)
            doc_buffer.seek(0)
            logging.info("Successfully generated traceability DOCX buffer.")
        except Exception as save_e:
            logging.error(f"Error saving traceability DOCX to buffer: {save_e}", exc_info=True)
            doc_buffer = BytesIO() # Return an empty buffer on error

        return doc_buffer

    def generate_health_snapshot_docx(self, project_name: str, snapshot_data: dict) -> BytesIO:
        """
        Generates the Project Health Snapshot .docx file, embedding Plotly charts
        into the styled template. (Phase 2 Final)
        """
        logging.info(f"Generating Project Pulse DOCX (Plotly) for project: {project_name}")
        document = self._get_styled_document()

        try:
            # --- Title ---
            timestamp_str = datetime.now(timezone.utc).isoformat()
            timestamp = format_timestamp_for_display(timestamp_str)

            try:
                document.add_paragraph(f"Project Pulse: {project_name}", style='Title')
                document.add_paragraph(f"Generated on: {timestamp}", style='Normal')
            except KeyError:
                logging.warning("Title or Normal style not found. Using fallbacks.")
                document.add_heading(f"Project Pulse: {project_name}", level=1)
                document.add_paragraph(f"Generated on: {timestamp}")

            document.add_paragraph() # Spacer

            # --- Data Extraction ---
            backlog_summary = snapshot_data.get('backlog_summary', {})
            test_summary = snapshot_data.get('test_summary', {})

            # --- 1. Backlog Completion Chart (Plotly) ---
            try:
                try:
                    document.add_paragraph("Backlog Completion Status", style='Heading 2')
                except KeyError:
                    document.add_heading("Backlog Completion Status", level=2)

                if backlog_summary:
                    labels = list(backlog_summary.keys())
                    values = list(backlog_summary.values())
                    fig = go.Figure(data=[go.Pie(labels=labels, values=values, textinfo='label+percent', hole=.3)])
                    fig.update_layout(title_text='Backlog Items by Status', showlegend=False)
                    image_bytes_io = generate_plotly_png(fig)
                    document.add_picture(image_bytes_io, width=Inches(6.0))
                else:
                    document.add_paragraph("No backlog status data available.", style='Normal')
            except Exception as e:
                logging.error(f"Failed to render Plotly backlog chart: {e}")
                document.add_paragraph(f"[Error rendering backlog chart: {e}]", style='Normal')

            document.add_paragraph() # Spacer

            # --- 2. Code Quality Chart (Plotly) ---
            try:
                try:
                    document.add_paragraph("Component Unit Test Status", style='Heading 2')
                except KeyError:
                    document.add_heading("Component Unit Test Status", level=2)

                if test_summary:
                    labels = list(test_summary.keys())
                    values = list(test_summary.values())
                    status_colors = {'PASSED': 'green', 'FAILED': 'red', 'NOT_TESTED': 'gray'}
                    colors = [status_colors.get(label, 'blue') for label in labels]
                    fig = go.Figure(data=[go.Bar(x=labels, y=values, marker_color=colors)])
                    fig.update_layout(title_text='Code Components by Test Status')
                    image_bytes_io = generate_plotly_png(fig)
                    document.add_picture(image_bytes_io, width=Inches(6.0))
                else:
                    document.add_paragraph("No component test data available.", style='Normal')
            except Exception as e:
                logging.error(f"Failed to render Plotly quality chart: {e}")
                document.add_paragraph(f"[Error rendering quality chart: {e}]", style='Normal')

        except Exception as e:
            logging.error(f"Error during health snapshot docx generation: {e}")
            try:
                document.add_paragraph(f"An unexpected error occurred: {e}", style='Normal')
            except KeyError:
                document.add_paragraph(f"An unexpected error occurred: {e}")

        # --- Save to in-memory buffer ---
        doc_buffer = BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer

    def generate_traceability_xlsx(self, trace_data: list, project_name: str) -> BytesIO:
        """
        Generates a formatted .xlsx file for the Requirements Traceability report.
        """
        logging.info(f"Generating XLSX for traceability report for project: {project_name}")
        wb = Workbook()
        ws = wb.active
        ws.title = "Traceability Matrix"

        # --- Headers ---
        headers = ['Backlog Item (#)', 'Title', 'Status', 'Artifact Path', 'Artifact Name']
        ws.append(headers)

        # Apply bold font to header row
        for cell in ws["1:1"]:
            cell.font = Font(bold=True)

        # --- Data ---
        if not trace_data:
            ws.append(["No traceability data available for this project."])
        else:
            for item in trace_data:
                ws.append([
                    item.get("backlog_id", "N/A"),
                    item.get("backlog_title", "N/A"),
                    item.get("backlog_status", "N/A"),
                    item.get("artifact_path", "N/A"),
                    item.get("artifact_name", "N/A")
                ])

        # Adjust column widths
        ws.column_dimensions['A'].width = 18
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 60
        ws.column_dimensions['E'].width = 30

        # Save to in-memory buffer
        xlsx_buffer = BytesIO()
        wb.save(xlsx_buffer)
        xlsx_buffer.seek(0)
        return xlsx_buffer

    def generate_sprint_deliverables_xlsx(self, sprint_id: str, report_data: list) -> BytesIO:
        """
        Generates an .xlsx file listing backlog items and linked artifacts for a sprint.
        """
        logging.info(f"ReportGenerator: Generating sprint deliverables XLSX for sprint {sprint_id}")
        wb = Workbook()
        ws = wb.active
        ws.title = "Sprint Deliverables"

        headers = ['Backlog Item (#)', 'Title', 'Status', 'Artifact Path', 'Artifact Name']
        ws.append(headers)
        for cell in ws["1:1"]:
            cell.font = Font(bold=True)

        if not report_data:
            ws.append(["No deliverables found or tracked for this sprint."])
        else:
            for item in report_data:
                ws.append([
                    item.get("backlog_id", "N/A"),
                    item.get("backlog_title", "N/A"),
                    item.get("backlog_status", "N/A"),
                    item.get("artifact_path", "N/A"),
                    item.get("artifact_name", "N/A")
                ])

        # Adjust column widths
        ws.column_dimensions['A'].width = 18
        ws.column_dimensions['B'].width = 50
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 60
        ws.column_dimensions['E'].width = 30

        xlsx_buffer = BytesIO()
        wb.save(xlsx_buffer)
        xlsx_buffer.seek(0)
        return xlsx_buffer

    def generate_burndown_chart_image(self, burndown_data: dict) -> BytesIO:
        """
        Generates a PNG image of the complexity point burndown chart using Plotly.
        (Placeholder implementation - needs refinement based on actual data structure)
        """
        logging.info(f"ReportGenerator: Generating Plotly burndown chart for sprint {burndown_data.get('sprint_id')}")

        try:
            # --- Placeholder Chart Logic ---
            # Assumes burndown_data has keys like 'total', 'remaining_per_step': [total, points_step1, points_step2,...]
            total_points = burndown_data.get('total', 0)
            # For this example, let's simulate a simple burndown
            steps = list(range(5)) # Simulate 5 tasks/days
            remaining = [total_points, total_points*0.8, total_points*0.5, total_points*0.2, 0] # Example data
            ideal_line = [total_points * (1 - (i / (len(steps)-1 if len(steps) > 1 else 1))) for i in steps]

            fig = go.Figure()

            # Add Ideal Burndown line
            fig.add_trace(go.Scatter(
                x=steps,
                y=ideal_line,
                mode='lines',
                name='Ideal Burndown',
                line=dict(color='gray', dash='dash')
            ))

            # Add Remaining Complexity line
            fig.add_trace(go.Scatter(
                x=steps,
                y=remaining,
                mode='lines+markers',
                name='Remaining Complexity',
                line=dict(color='#007ACC'),
                marker=dict(color='#007ACC')
            ))

            fig.update_layout(
                title=f"Sprint Burndown ({burndown_data.get('sprint_id', 'N/A')})",
                xaxis_title="Tasks Completed (Sequence)",
                yaxis_title="Remaining Complexity Points",
                yaxis_range=[0, total_points * 1.05], # Set y-axis to start at 0
                template="plotly_dark"
            )

            img_buffer = generate_plotly_png(fig)
            return img_buffer
        except Exception as e:
            logging.error(f"Failed to generate burndown chart image: {e}", exc_info=True)
            # Re-raise the exception to be handled by the caller
            raise
        # --- End Placeholder ---

    def generate_cfd_chart_image(self, cfd_data: dict) -> BytesIO:
        """
        Generates a PNG image of the Cumulative Flow Diagram using Plotly.
        (Placeholder implementation - needs historical data)
        """
        logging.info("ReportGenerator: Generating Plotly CFD chart image.")

        try:
            # --- Placeholder Chart Logic ---
            # Assumes cfd_data contains historical counts per status per day/event
            # Example data structure: {'dates': [d1, d2, d3], 'TO_DO': [10, 8, 5], 'IN_PROGRESS': [0, 2, 3], 'COMPLETED': [0, 0, 2]}
            dates = ['Day 1', 'Day 2', 'Day 3', 'Day 4', 'Day 5']
            statuses = {
                'TO_DO': [10, 8, 5, 5, 2],
                'ANALYZED': [0, 1, 2, 1, 1],
                'IN_PROGRESS': [0, 2, 3, 4, 3],
                'COMPLETED': [0, 0, 2, 2, 6]
            }

            fig = go.Figure()

            # Plotly's stackplot is created by adding traces with 'tonexty' fill
            last_y = [0] * len(dates)
            for status_name, values in statuses.items():
                current_y = [last + val for last, val in zip(last_y, values)]
                fig.add_trace(go.Scatter(
                    x=dates,
                    y=current_y,
                    fill='tonexty',
                    mode='none',
                    name=status_name
                ))
                last_y = current_y # Update the baseline for the next layer

            fig.update_layout(
                title="Workflow Efficiency (Cumulative Flow)",
                xaxis_title="Time",
                yaxis_title="Number of Items",
                template="plotly_dark"
            )

            img_buffer = generate_plotly_png(fig)
            return img_buffer
        except Exception as e:
            logging.error(f"Failed to generate CFD chart image: {e}", exc_info=True)
            raise
        # --- End Placeholder ---

    def generate_quality_trend_chart_image(self, trend_data: dict) -> BytesIO:
        """
        Generates a PNG image of the Code Quality Trend chart using Plotly.
        (Placeholder implementation - needs historical data)
        """
        logging.info("ReportGenerator: Generating Plotly code quality trend chart image.")

        try:
            # --- Placeholder Chart Logic ---
            # Assumes trend_data contains snapshots over time/sprints
            # Example: {'sprints': ['S1', 'S2', 'S3'], 'PASSED': [5, 8, 12], 'FAILED': [2, 1, 0], 'NOT_TESTED': [3, 1, 0]}
            sprints = ['Sprint 1', 'Sprint 2', 'Sprint 3']
            passed = [5, 8, 12]
            failed = [2, 1, 0]
            not_tested = [3, 1, 0]

            fig = go.Figure()

            fig.add_trace(go.Scatter(
                x=sprints,
                y=passed,
                mode='lines+markers',
                name='Tests Passing',
                line=dict(color='#6A8759') # Green
            ))

            fig.add_trace(go.Scatter(
                x=sprints,
                y=failed,
                mode='lines+markers',
                name='Tests Failing',
                line=dict(color='#CC7832') # Red/Orange
            ))

            # Optionally plot not_tested if useful
            fig.add_trace(go.Scatter(
                x=sprints,
                y=not_tested,
                mode='lines+markers',
                name='Not Tested',
                line=dict(color='#888888', dash='dash') # Gray
            ))

            fig.update_layout(
                title="Code Quality Trend (Unit Test Status)",
                xaxis_title="Sprint",
                yaxis_title="Number of Components",
                yaxis_range=[0, max(passed + failed + not_tested) * 1.05], # Start y-axis at 0
                template="plotly_dark",
                legend_title_text="Test Status"
            )

            img_buffer = generate_plotly_png(fig)
            return img_buffer
        except Exception as e:
            logging.error(f"Failed to generate code quality trend chart image: {e}", exc_info=True)
            raise
        # --- End Placeholder ---

    def generate_ai_assistance_report(self, assistance_data: dict) -> BytesIO:
        """
        Formats the AI Assistance Rate data into a .docx report.
        """
        logging.info("ReportGenerator: Generating AI assistance rate report text.")
        total_sprints = assistance_data.get('total_sprints_analyzed', 0)
        total_escalations = assistance_data.get('total_escalations', 0)
        avg_per_sprint = assistance_data.get('average_escalations_per_sprint', 0)

        if total_sprints == 0:
            report_text = "No sprint data available to calculate AI Assistance Rate."
        else:
            report_lines = [
                "## AI Assistance Rate Summary\n",
                f"- **Total Sprints Analyzed:** {total_sprints}",
                f"- **Total Debug Escalations to PM:** {total_escalations}",
                f"- **Average Escalations per Sprint:** {avg_per_sprint:.2f}\n",
                "_Lower rates indicate higher AI reliability and autonomy._"
            ]
            report_text = "\n".join(report_lines)

        # Call the existing docx generator to wrap the text content
        docx_bytes_io = self.generate_text_document_docx(
            title="AI Assistance Rate Report",
            content=report_text,
            is_code=False
        )
        return docx_bytes_io

    def generate_backlog_xlsx(self, backlog_data: list) -> BytesIO:
        """
        Generates an .xlsx file for the full or filtered project backlog.
        Accepts a flat list containing hierarchical IDs.
        """
        logging.info("ReportGenerator: Generating backlog XLSX.")
        flat_list = []

        # Check if input is hierarchical or already flat list from filtered query
        if backlog_data and 'hierarchical_id' in backlog_data[0]:
            # Input is already flat (or we're treating it as such for filtered export)
            for item in backlog_data:
                timestamp_str = item.get('last_modified_timestamp') or item.get('creation_timestamp')
                formatted_date = format_timestamp_for_display(timestamp_str) if timestamp_str else "N/A"
                record = {
                    '#': item.get('hierarchical_id', f"CR-{item.get('cr_id', 'N/A')}"),
                    'Title': item.get('title', 'N/A'),
                    'Type': item.get('request_type', 'N/A').replace('_', ' ').title(),
                    'Status': item.get('status', 'N/A'),
                    'Priority/Severity': item.get('priority') or item.get('impact_rating') or '',
                    'Complexity': item.get('complexity', ''),
                    'Last Modified': formatted_date
                }
                flat_list.append(record)
        else:
            # Input is hierarchical, need to flatten (original logic)
            def flatten_hierarchy(items, prefix=""):
                # ...(same flattening logic as before)...
                for i, item in enumerate(items, 1):
                    current_prefix = f"{prefix}{i}"
                    timestamp_str = item.get('last_modified_timestamp') or item.get('creation_timestamp')
                    formatted_date = format_timestamp_for_display(timestamp_str) if timestamp_str else "N/A"
                    record = {
                        '#': current_prefix,
                        'Title': item.get('title', 'N/A'),
                        'Type': item.get('request_type', 'N/A').replace('_', ' ').title(),
                        'Status': item.get('status', 'N/A'),
                        'Priority/Severity': item.get('priority') or item.get('impact_rating') or '',
                        'Complexity': item.get('complexity', ''),
                        'Last Modified': formatted_date
                    }
                    flat_list.append(record)
                    if "features" in item: flatten_hierarchy(item["features"], prefix=f"{current_prefix}.")
                    if "user_stories" in item: flatten_hierarchy(item["user_stories"], prefix=f"{current_prefix}.")
            flatten_hierarchy(backlog_data)

        if not flat_list:
            flat_list.append({'#': "No items match filter criteria.", 'Title': '', 'Type': '', 'Status': '', 'Priority/Severity': '', 'Complexity': '', 'Last Modified': ''})

        df = pd.DataFrame(flat_list)
        output_buffer = BytesIO()
        with pd.ExcelWriter(output_buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Project Backlog')
        output_buffer.seek(0)
        return output_buffer