# agents/agent_project_bootstrap.py

"""
This module contains the ProjectBootstrapAgent class.
"""

import logging
from typing import List, Tuple, Optional
from asdf_db_manager import ASDFDBManager
from pathlib import Path
# Note: pypdf is now imported inside the method to prevent startup hangs.

class ProjectBootstrapAgent:
    """
    Processes uploaded specification documents for the target application.
    """
    def __init__(self, db_manager: ASDFDBManager):
        """
        Initializes the ProjectBootstrapAgent.
        """
        if not db_manager:
            raise ValueError("db_manager is required for the ProjectBootstrapAgent.")
        self.db_manager = db_manager

    def extract_text_from_files(self, uploaded_files: List) -> Tuple[Optional[str], List[str], Optional[str]]:
        """
        Extracts text content from a list of uploaded files and checks size limits.
        """
        import docx
        all_text = []
        messages = []
        for doc in uploaded_files:
            try:
                if doc.name.endswith('.docx'):
                    document = docx.Document(doc)
                    for para in document.paragraphs:
                        all_text.append(para.text)
                elif doc.name.endswith('.pdf'):
                    from pypdf import PdfReader
                    try:
                        pdf_reader = PdfReader(doc)
                        for page in pdf_reader.pages:
                            all_text.append(page.extract_text() or "")
                    except Exception as pdf_e:
                        messages.append(f"Warning: Could not read PDF file '{doc.name}'. It may be corrupted or image-based. Error: {pdf_e}")
                elif doc.name.endswith(('.txt', '.md')):
                    all_text.append(doc.getvalue().decode("utf-8"))
                else:
                    messages.append(f"Warning: Unsupported file type '{doc.name}' was skipped.")
            except Exception as e:
                messages.append(f"Error: Could not process file '{doc.name}'. Reason: {e}")

        concatenated_text = "\n\n---\n\n".join(all_text)

        try:
            limit_str = self.db_manager.get_config_value("CONTEXT_WINDOW_CHAR_LIMIT") or "2000000"
            spec_max_char_limit = int(limit_str)
        except Exception as e:
            logging.error(f"Could not read CONTEXT_WINDOW_CHAR_LIMIT from DB, using default. Error: {e}")
            spec_max_char_limit = 128000

        if len(concatenated_text) > spec_max_char_limit:
            error_message = (
                "The provided specification is too large for a single project based on the current LLM's context limit. "
                "Please divide it into smaller, more focused sub-projects or select an LLM with a larger context window in Settings."
            )
            messages.append(f"Error: Specification character count ({len(concatenated_text):,}) exceeds the active limit of {spec_max_char_limit:,}.")
            return None, messages, error_message

        return concatenated_text, messages, None

    def extract_text_from_file_paths(self, file_paths: list[str]) -> tuple[str | None, list[str], str | None]:
        """
        Extracts text content from a list of local file paths and checks size limits.
        """
        import docx
        all_text = []
        messages = []
        for path_str in file_paths:
            try:
                path = Path(path_str)
                if not path.exists():
                    messages.append(f"Warning: File not found at path '{path_str}' and was skipped.")
                    continue

                if path.name.endswith('.docx'):
                    document = docx.Document(path)
                    for para in document.paragraphs:
                        all_text.append(para.text)
                elif path.name.endswith('.pdf'):
                    from pypdf import PdfReader
                    try:
                        with open(path, "rb") as f:
                            pdf_reader = PdfReader(f)
                            for page in pdf_reader.pages:
                                all_text.append(page.extract_text() or "")
                    except Exception as pdf_e:
                        messages.append(f"Warning: Could not read PDF file '{path.name}'. It may be corrupted or image-based. Error: {pdf_e}")
                elif path.name.endswith(('.txt', '.md')):
                    all_text.append(path.read_text(encoding="utf-8"))
                else:
                    messages.append(f"Warning: Unsupported file type '{path.name}' was skipped.")
            except Exception as e:
                messages.append(f"Error: Could not process file '{path_str}'. Reason: {e}")

        concatenated_text = "\n\n---\n\n".join(all_text)

        try:
            limit_str = self.db_manager.get_config_value("CONTEXT_WINDOW_CHAR_LIMIT") or "2000000"
            spec_max_char_limit = int(limit_str)
        except Exception as e:
            logging.error(f"Could not read CONTEXT_WINDOW_CHAR_LIMIT from DB, using default. Error: {e}")
            spec_max_char_limit = 128000

        if len(concatenated_text) > spec_max_char_limit:
            error_message = (
                "The provided specification is too large for a single project based on the current LLM's context limit. "
                "Please divide it into smaller, more focused sub-projects or select an LLM with a larger context window in Settings."
            )
            messages.append(f"Error: Specification character count ({len(concatenated_text):,}) exceeds the active limit of {spec_max_char_limit:,}.")
            return None, messages, error_message

        return concatenated_text, messages, None