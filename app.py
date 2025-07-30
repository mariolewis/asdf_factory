import streamlit as st
from pathlib import Path
import time
import pandas as pd
from datetime import datetime
import logging
import docx
import json

# Imports from the project's root directory
from master_orchestrator import MasterOrchestrator
from agents.agent_environment_setup_app_target import EnvironmentSetupAgent_AppTarget
from agents.agent_project_bootstrap import ProjectBootstrapAgent
from agents.agent_spec_clarification import SpecClarificationAgent
from agents.agent_planning_app_target import PlanningAgent_AppTarget
from agents.agent_report_generator import ReportGeneratorAgent
from agents.agent_project_scoping import ProjectScopingAgent
from agents.agent_tech_stack_proposal import TechStackProposalAgent
from agents.agent_build_script_generator import BuildScriptGeneratorAgent
from agents.agent_verification_app_target import VerificationAgent_AppTarget
from agents.agent_coding_standard_app_target import CodingStandardAgent_AppTarget

# --- Page Configuration ---
st.set_page_config(
    page_title="ASDF",
    page_icon="🤖",
    layout="wide"
)

# --- Application State Management ---
db_dir = Path("data")
db_dir.mkdir(exist_ok=True)
db_path = db_dir / "asdf.db"

# Initialize session state variables
if 'orchestrator' not in st.session_state:
    st.session_state.orchestrator = MasterOrchestrator(db_path=str(db_path))
    with st.session_state.orchestrator.db_manager as db:
        api_key = db.get_config_value("LLM_API_KEY")
        if not api_key:
            st.warning("LLM API Key not found. Please go to the Settings page to configure it.")

if "messages" not in st.session_state:
    st.session_state.messages = []

if "api_key_input" not in st.session_state:
    st.session_state.api_key_input = ""


# --- Sidebar ---
with st.sidebar:
    st.markdown("## 🤖 Autonomous Software Development Factory")
    st.markdown("---")
    page = st.radio("Navigation", ["Project", "Documents", "Reports", "Settings"], label_visibility="collapsed")
    st.markdown("---")

    st.markdown("### Project Information")
    status_info = st.session_state.orchestrator.get_status()
    st.markdown(f"**Project ID:** {status_info.get('project_id', 'N/A')}")
    st.markdown(f"**Project Name:** {status_info.get('project_name', 'N/A')}")

    current_phase_enum = st.session_state.orchestrator.current_phase
    display_phase_name = st.session_state.orchestrator.PHASE_DISPLAY_NAMES.get(current_phase_enum, current_phase_enum.name)
    st.markdown(f"**Current Phase:** {display_phase_name}")

    st.markdown("---")
    st.markdown("### Actions")

    # Only show these buttons if a project is active
    if st.session_state.orchestrator.project_id:
        if st.button("✍️ Raise CR", use_container_width=True):
            st.session_state.orchestrator.handle_raise_cr_action()
            st.rerun()

        if st.button("🐞 Report Bug", use_container_width=True):
            st.session_state.orchestrator.handle_report_bug_action()
            st.rerun()

        # Disable the Implement button if Genesis is not complete
        genesis_complete = st.session_state.orchestrator.is_genesis_complete
        if st.button("🔁 Implement CR/Bug", use_container_width=True, disabled=not genesis_complete, help="Enabled after initial development is complete."):
            st.session_state.orchestrator.handle_view_cr_register_action()
            st.rerun()

    st.markdown("---")
    st.markdown("### Project Lifecycle")

    # --- CORRECTED & SIMPLIFIED RESUME LOGIC ---
    # The orchestrator now knows its own state on initialization. We just ask it.
    if st.session_state.orchestrator.resumable_state:
        if st.button("▶️ Resume Paused Project", use_container_width=True, type="primary"):
            if st.session_state.orchestrator.resume_project():
                st.rerun()
            else:
                st.error("Failed to resume project. Check logs for details.")

    # Show Stop button ONLY if a project is active and NOT in a resumable state.
    if st.session_state.orchestrator.project_id and not st.session_state.orchestrator.resumable_state:
        if st.button("⏹️ Stop & Export Active Project", use_container_width=True):
            st.session_state.show_export_confirmation = True

    if st.button("📂 Load Archived Project", use_container_width=True):
        st.session_state.orchestrator.set_phase("VIEWING_PROJECT_HISTORY")
        st.rerun()

    # The 'show_export_confirmation' logic remains the same as before
    if st.session_state.get("show_export_confirmation"):
        with st.form("export_form"):
            st.warning(
                "**Archive Project Confirmation**\n\n"
                "This will save all project data to an external archive file and clear the active session."
            )
            archive_name_input = st.text_input(
                "Enter a name for the archive file:",
                value=f"{st.session_state.orchestrator.project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}"
            )
            with st.session_state.orchestrator.db_manager as db:
                archive_path = db.get_config_value("DEFAULT_ARCHIVE_PATH") or "data/archives"

            submitted = st.form_submit_button("Confirm and Export")
            if submitted:
                if archive_name_input:
                    archive_file_path = st.session_state.orchestrator.stop_and_export_project(archive_path, archive_name_input)
                    if archive_file_path:
                        # --- CORRECTED: Perform a hard reset of the orchestrator state ---
                        st.session_state.last_action_success_message = f"Project archived to: `{archive_file_path}`"
                        # Re-initialize the orchestrator to ensure a clean state.
                        st.session_state.orchestrator = MasterOrchestrator(db_path=str(db_path))
                        st.session_state.show_export_confirmation = False
                        st.rerun()
                    else:
                        st.error("Failed to export project.")
                else:
                    st.error("Archive name cannot be empty.")

# --- Main Application UI ---
st.markdown("""
    <script>
        const streamlitDoc = window.parent.document;
        const observer = new MutationObserver(function(mutations) {
            mutations.forEach(function(mutation) {
                if (mutation.type === 'childList' && mutation.addedNodes.length) {
                    streamlitDoc.querySelector('.main').scrollTo(0, 0);
                }
            });
        });
        observer.observe(streamlitDoc.body, { childList: true, subtree: true });
    </script>
    """, unsafe_allow_html=True)

if page == "Project":
    if "last_action_success_message" in st.session_state:
        st.success(st.session_state.last_action_success_message)
        del st.session_state.last_action_success_message

    # Get the current phase name once at the top
    status_info = st.session_state.orchestrator.get_status()
    current_phase_name = status_info.get("current_phase")

    # --- CORRECTED LOGIC: Check for specific phases before checking for an active project ---
    if current_phase_name == "VIEWING_PROJECT_HISTORY":
        st.header("Load an Archived Project")

        project_history = st.session_state.orchestrator.get_project_history()

        if not project_history:
            st.warning("There are no archived projects in the history.")
        else:
            history_data = [
                {
                    "ID": row['history_id'],
                    "Project Name": row['project_name'],
                    "Project ID": row['project_id'],
                    "Archived On": row['last_stop_timestamp'],
                    "Archive Path": row['archive_file_path']
                }
                for row in project_history
            ]
            df = pd.DataFrame(history_data)
            st.dataframe(df, use_container_width=True, hide_index=True, column_config={"ID": "Select"})

            history_ids = [row['history_id'] for row in project_history]
            selected_id_str = st.selectbox("Select a Project ID to action:", options=[""] + [str(i) for i in history_ids])

            if selected_id_str:
                selected_id = int(selected_id_str)
                selected_project_record = next((p for p in project_history if p['history_id'] == selected_id), None)

                is_active_project = (st.session_state.orchestrator.project_id == selected_project_record['project_id'])

                col1, col2, _ = st.columns([1, 1, 4])
                with col1:
                    if st.button("📂 Load Selected Project", use_container_width=True, type="primary"):
                        with st.spinner("Loading project data and running pre-flight checks..."):
                            # Save the selected ID to session state so it's available on the next screen
                            st.session_state.selected_history_id_for_action = selected_id
                            st.session_state.orchestrator.load_archived_project(selected_id)
                            st.rerun()
                with col2:
                    delete_button = st.button("🗑️ Delete Selected Project", use_container_width=True, disabled=is_active_project, help="You cannot delete the currently active project.")

                if delete_button:
                    with st.popover("Confirm Deletion", use_container_width=True):
                        st.warning(f"**Are you sure you want to permanently delete project '{selected_project_record['project_name']}' (ID: {selected_id})?**")
                        st.error("This action cannot be undone and will delete the project history record and its associated archive files from the disk.")
                        if st.button("Yes, permanently delete this project", type="primary", use_container_width=True):
                            success, message = st.session_state.orchestrator.delete_archived_project(selected_id)
                            if success:
                                st.toast(f"✅ {message}")
                            else:
                                st.error(message)
                            time.sleep(2)
                            st.rerun()

        st.divider()
        if st.button("⬅️ Back to Main Page"):
            # Perform a hard reset to ensure a clean return to the idle state
            st.session_state.orchestrator = MasterOrchestrator(db_path=str(db_path))
            st.rerun()

    elif current_phase_name == "AWAITING_PREFLIGHT_RESOLUTION":
        st.header("Pre-flight Check Resolution")
        result = st.session_state.orchestrator.preflight_check_result
        if not result:
            st.error("Error: Pre-flight check result not found.")
            if st.button("Go Back"):
                st.session_state.orchestrator.set_phase("IDLE")
                st.rerun()
        else:
            status = result.get("status")
            message = result.get("message")
            if status == "ALL_PASS":
                st.success(message)
                st.info("The project environment is valid and clean. You can now proceed with the project.")

                resume_phase = st.session_state.orchestrator.resume_phase_after_load
                if resume_phase:
                    phase_display_name = st.session_state.orchestrator.PHASE_DISPLAY_NAMES.get(resume_phase, resume_phase.name.replace("_", " ").title())

                    col1, col2, _ = st.columns([1.5, 2, 3])
                    with col1:
                        if st.button(f"▶️ Proceed to {phase_display_name}", type="primary", use_container_width=True):
                            # If resuming to GENESIS, we must load the dev plan into active memory.
                            if resume_phase.name == "GENESIS":
                                with st.spinner("Loading development plan..."):
                                    with st.session_state.orchestrator.db_manager as db:
                                        project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                        if project_details and project_details['development_plan_text']:
                                            full_plan_data = json.loads(project_details['development_plan_text'])
                                            dev_plan_list = full_plan_data.get("development_plan")
                                            if dev_plan_list:
                                                st.session_state.orchestrator.load_development_plan(json.dumps(dev_plan_list))

                            st.session_state.orchestrator.set_phase(resume_phase.name)
                            st.session_state.orchestrator.resume_phase_after_load = None # Clean up
                            st.rerun()
                    with col2:
                        if st.button("⬅️ Back to Project List", use_container_width=True):
                            st.session_state.orchestrator = MasterOrchestrator(db_path=str(db_path))
                            st.session_state.orchestrator.set_phase("VIEWING_PROJECT_HISTORY")
                            st.rerun()
                else:
                    st.error("Could not determine the resume phase. Please go back and try loading the project again.")
                    if st.button("⬅️ Back to Project List"):
                        st.session_state.orchestrator = MasterOrchestrator(db_path=str(db_path))
                        st.session_state.orchestrator.set_phase("VIEWING_PROJECT_HISTORY")
                        st.rerun()
            elif status in ["PATH_NOT_FOUND", "GIT_MISSING", "ERROR"]:
                st.error(message)
                st.warning("The factory cannot proceed until this environmental issue is resolved.")
                if st.button("Go to Environment Setup to Resolve"):
                    # This is a fatal error, so we create a new project with the same name.
                    # This will cleanly guide the user through the setup flow again.
                    project_name = selected_project_record['project_name'] if selected_project_record else "New Project"
                    st.session_state.orchestrator.start_new_project(project_name)
                    st.rerun()
            elif status == "STATE_DRIFT":
                st.warning(message)
                col1, col2, _ = st.columns([1.5, 1.5, 3])
                with col1:
                    if st.button("I will resolve this manually", help="Use your own tools (e.g., git commit, git stash) to clean the repository, then load the project again from the main menu."):
                        st.session_state.orchestrator.project_id = None
                        st.session_state.orchestrator.set_phase("IDLE")
                        st.rerun()
                with col2:
                    with st.popover("Expert Option: Discard"):
                        st.markdown("⚠️ **This will permanently delete all uncommitted changes in your local repository. This cannot be undone.**")
                        if st.button("Confirm & Discard All Changes", type="primary"):
                            with st.spinner("Resetting repository and re-checking..."):
                                project_id_to_reset = st.session_state.orchestrator.project_id
                                st.session_state.orchestrator.handle_discard_changes(history_id=st.session_state.selected_history_id_for_action)
                            st.rerun()

            elif current_phase_name == "AWAITING_REASSESSMENT_CONFIRMATION":
                st.header("Mid-Project Re-assessment Required")

                # This UI has two stages: initial confirmation, and final decision after report.
                reassessment_data = st.session_state.orchestrator.task_awaiting_approval or {}
                reassessment_result = reassessment_data.get("reassessment_result")

                if not reassessment_result:
                    # Stage 1: Initial Confirmation
                    st.warning(
                        "You have selected an LLM Provider with a smaller default context window than the one currently active. "
                        "This may increase the risk of failure for complex tasks on the remaining work in this project."
                    )
                    st.markdown("It is recommended to run a re-assessment on the project's remaining scope before proceeding.")
                    st.divider()

                    col1, col2, _ = st.columns([1.5, 2, 3])
                    with col1:
                        if st.button("▶️ Proceed with Re-assessment", type="primary", use_container_width=True):
                            with st.spinner("Analyzing remaining project scope..."):
                                st.session_state.orchestrator.run_mid_project_reassessment()
                            st.rerun()
                    with col2:
                        if st.button("❌ Cancel and Revert LLM Choice", use_container_width=True):
                            # Clean up state and return to GENESIS
                            keys_to_clear = ['reassessment_required', 'pending_llm_provider', 'previous_llm_provider']
                            for key in keys_to_clear:
                                if key in st.session_state: del st.session_state[key]
                            st.session_state.orchestrator.set_phase("GENESIS")
                            st.toast("LLM provider change was cancelled.")
                            st.rerun()
                else:
                    # Stage 2: Review Report and Make Final Decision
                    st.subheader("Re-assessment Report for Remaining Work")
                    if "error" in reassessment_result:
                        st.error(f"Could not generate re-assessment report: {reassessment_result['error']}")
                    else:
                        # Display a simplified version of the assessment report
                        risk = reassessment_result.get('risk_assessment', {})
                        st.metric("Overall Risk Level for Remaining Work", risk.get('overall_risk_level', 'N/A'))
                        st.write("**Risk Summary:**")
                        st.write(risk.get('summary', 'No summary provided.'))

                    st.divider()
                    st.markdown(f"Do you want to finalize the switch to **{st.session_state.get('pending_llm_provider')}** or revert to **{st.session_state.get('previous_llm_provider')}**?")

                    col1, col2, _ = st.columns([2, 2, 3])
                    with col1:
                        if st.button(f"✅ Continue with {st.session_state.get('pending_llm_provider')}", type="primary", use_container_width=True):
                            success, message = st.session_state.orchestrator.commit_pending_llm_change(st.session_state.pending_llm_provider)
                            if success:
                                st.toast(message, icon="✅")
                            else:
                                st.error(message)
                            keys_to_clear = ['reassessment_required', 'pending_llm_provider', 'previous_llm_provider']
                            for key in keys_to_clear:
                                if key in st.session_state: del st.session_state[key]
                            st.session_state.orchestrator.set_phase("GENESIS")
                            st.rerun()
                    with col2:
                        if st.button(f"❌ Revert to {st.session_state.get('previous_llm_provider')}", use_container_width=True):
                            keys_to_clear = ['reassessment_required', 'pending_llm_provider', 'previous_llm_provider']
                            for key in keys_to_clear:
                                if key in st.session_state: del st.session_state[key]
                            st.session_state.orchestrator.set_phase("GENESIS")
                            st.toast("LLM provider change was reverted.")
                            st.rerun()

    elif not st.session_state.orchestrator.project_id:
        st.subheader("Start a New Project")
        st.info(
            "**Note:** If another project is currently active in the factory, "
            "it will be automatically saved and archived before the new project begins."
        )
        project_name_input = st.text_input("Enter a name for your new project:")
        if st.button("Start New Project"):
            if project_name_input:
                st.session_state.orchestrator.start_new_project(project_name_input)
                st.rerun()
            else:
                st.error("Please enter a project name.")
    else:
        # This 'else' block now correctly handles all other phases for an active project.
        if current_phase_name == "AWAITING_UX_UI_PHASE_DECISION":
                st.header("New Project Intake")
                st.markdown("To help guide the development process, please provide the initial project brief. You can either type a brief description or upload an existing document.")

                tab1, tab2 = st.tabs(["Enter Brief Description", "Upload Brief Document"])

                with tab1:
                    if 'project_brief_input' not in st.session_state:
                        st.session_state.project_brief_input = ""

                    brief_desc = st.text_area(
                        "Project Brief:",
                        value=st.session_state.project_brief_input,
                        height=150,
                        key="project_brief_input"
                    )
                    if st.button("Analyze Text Brief", type="primary"):
                        if brief_desc.strip():
                            with st.spinner("Analyzing brief..."):
                                st.session_state.orchestrator.handle_ux_ui_brief_submission(brief_desc)
                            st.rerun()
                        else:
                            st.warning("Please provide a brief description before analyzing.")

                with tab2:
                    uploaded_brief = st.file_uploader(
                        "Upload Brief Document",
                        type=['txt', 'md', 'docx'],
                        label_visibility="collapsed"
                    )
                    if st.button("Analyze Uploaded Brief", type="primary"):
                        if uploaded_brief is not None:
                            with st.spinner("Analyzing brief..."):
                                st.session_state.orchestrator.handle_ux_ui_brief_submission(uploaded_brief)
                            st.rerun()
                        else:
                            st.warning("Please upload a document before analyzing.")

        elif current_phase_name == "AWAITING_UX_UI_RECOMMENDATION_CONFIRMATION":
                st.header("UX/UI Phase Recommendation")

                approval_task = st.session_state.orchestrator.task_awaiting_approval or {}

                if "analysis_error" in approval_task:
                    st.error(f"An error occurred during the initial analysis: {approval_task['analysis_error']}")
                    st.warning("Please choose how to proceed manually.")
                else:
                    analysis = approval_task.get("analysis", {})
                    recommendation = analysis.get("ux_phase_necessity", "Optional")
                    justification = analysis.get("justification", "No justification was provided.")

                    st.info(f"**AI Recommendation: {recommendation}**")
                    st.write(justification)

                    personas = analysis.get("inferred_personas", [])
                    if personas:
                        st.write("**Inferred User Personas:**")
                        st.markdown("- " + "\n- ".join(personas))

                st.divider()
                st.markdown("How would you like to proceed?")

                col1, col2, _ = st.columns([1.5, 2, 3])
                with col1:
                    # The recommended option is highlighted as the primary button
                    is_primary = recommendation in ["Recommended", "Optional"]
                    if st.button("Start UX/UI Design Phase", type="primary" if is_primary else "secondary", use_container_width=True):
                        st.session_state.orchestrator.handle_ux_ui_phase_decision("START_UX_UI_PHASE")
                        st.rerun()

                with col2:
                    is_primary_skip = recommendation == "Not Recommended"
                    if st.button("Skip to Application Spec", type="primary" if is_primary_skip else "secondary", use_container_width=True):
                        st.session_state.orchestrator.handle_ux_ui_phase_decision("SKIP_TO_SPEC")
                        st.rerun()

        elif current_phase_name == "UX_UI_DESIGN":
                st.header("User Experience & Interface Design")

                # Initialize state for this multi-step phase
                if 'ux_design_step' not in st.session_state:
                    st.session_state.ux_design_step = 'confirm_personas'

                if st.session_state.ux_design_step == 'confirm_personas':
                    st.subheader("Step 1: Confirm User Personas")
                    st.markdown("The AI has inferred the following user personas/roles from your project brief. Please review, edit if necessary, and confirm.")

                    # Load inferred personas into an editable text area
                    if 'personas_text' not in st.session_state:
                        analysis = st.session_state.orchestrator.task_awaiting_approval.get("analysis", {})
                        initial_personas = analysis.get("inferred_personas", ["Default User"])
                        st.session_state.personas_text = "\n".join(initial_personas)

                    edited_personas = st.text_area(
                        "User Personas (one per line):",
                        value=st.session_state.personas_text,
                        height=150,
                        key="personas_text_area"
                    )

                    if st.button("Confirm Personas & Generate User Journeys", type="primary"):
                        if edited_personas.strip():
                            persona_list = [p.strip() for p in edited_personas.strip().split('\n') if p.strip()]
                            with st.spinner("Generating core user journeys based on personas..."):
                                # We will create this backend method in the next step.
                                st.session_state.orchestrator.handle_ux_persona_confirmation(persona_list)
                            st.session_state.ux_design_step = 'review_user_journeys'
                            st.rerun()
                        else:
                            st.warning("Please define at least one user persona.")

                elif st.session_state.ux_design_step == 'review_user_journeys':
                    st.subheader("Step 2: Review User Journeys")
                    st.markdown("The AI has generated a list of core user journeys based on the personas. Please review, edit if necessary, and confirm.")

                    # Load generated journeys into an editable text area
                    if 'journeys_text' not in st.session_state:
                        journeys = st.session_state.orchestrator.active_ux_spec.get('generated_user_journeys', '1. Default user journey.')
                        st.session_state.journeys_text = journeys

                    edited_journeys = st.text_area(
                        "Core User Journeys:",
                        value=st.session_state.journeys_text,
                        height=200,
                        key="journeys_text_area"
                    )

                    if st.button("Confirm Journeys & Identify Screens", type="primary"):
                        if edited_journeys.strip():
                            with st.spinner("Analyzing journeys to identify application screens..."):
                                st.session_state.orchestrator.handle_ux_journey_confirmation(edited_journeys)
                            st.session_state.ux_design_step = 'review_screens'
                            st.rerun()
                        else:
                            st.warning("Please define at least one user journey.")

                elif st.session_state.ux_design_step == 'review_screens':
                    st.subheader("Step 3: Review Application Screens")
                    st.markdown("Based on the user journeys, the AI has identified the following necessary screens/views for the application. Please review, edit, and confirm this list.")

                    # Load identified screens into an editable text area
                    if 'screens_text' not in st.session_state:
                        screens = st.session_state.orchestrator.active_ux_spec.get('identified_screens', '1. Main Screen')
                        st.session_state.screens_text = screens

                    edited_screens = st.text_area(
                        "Application Screens/Views (one per line):",
                        value=st.session_state.screens_text,
                        height=200,
                        key="screens_text_area"
                    )

                    if st.button("Confirm Screens & Begin Detailed Design", type="primary"):
                        if edited_screens.strip():
                            st.session_state.orchestrator.handle_ux_screen_confirmation(edited_screens)
                            st.session_state.ux_design_step = 'detailed_screen_design'
                            st.rerun()
                        else:
                            st.warning("Please define at least one screen.")

                elif st.session_state.ux_design_step == 'detailed_screen_design':
                    st.subheader("Step 4: Detailed Screen Design")

                    # Get the list of screens to design
                    screens_str = st.session_state.orchestrator.active_ux_spec.get('confirmed_screens_text', '1. Main Screen')
                    screen_list = [line.strip() for line in screens_str.split('\n') if line.strip()]

                    # Get the current screen index
                    cursor = st.session_state.orchestrator.active_ux_spec.get('screen_design_cursor', 0)

                    if cursor >= len(screen_list):
                        st.success("All screen blueprints have been designed.")
                        st.info("The next step is to define the overall theme and style guide for the application.")
                        if st.button("Proceed to Style Guide", type="primary"):
                            st.session_state.ux_design_step = 'define_style_guide'
                            st.rerun()
                        st.stop()

                    current_screen_name = screen_list[cursor].split('. ', 1)[-1] # Clean the "1. " prefix

                    st.markdown(f"#### Designing Screen {cursor + 1} of {len(screen_list)}: **{current_screen_name}**")

                    st.markdown("Describe the components, layout, and functionality of this screen in the text area below.")

                    pm_description = st.text_area(
                        f"Description for {current_screen_name}:",
                        height=200,
                        key=f"desc_{current_screen_name}"
                    )

                    if st.button(f"Generate Blueprint for {current_screen_name}"):
                        if pm_description.strip():
                            with st.spinner(f"Generating JSON blueprint for {current_screen_name}..."):
                                st.session_state.orchestrator.handle_screen_design_submission(current_screen_name, pm_description)
                            st.rerun()
                        else:
                            st.warning("Please provide a description for the screen.")

                    # Display the generated blueprint for review
                    blueprints = st.session_state.orchestrator.active_ux_spec.get('screen_blueprints', {})
                    if current_screen_name in blueprints:
                        st.markdown("**Generated Blueprint (JSON):**")
                        st.json(blueprints[current_screen_name])

                    # Navigation and finalization buttons
                    st.divider()
                    col1, col2, col3 = st.columns([1.5, 1.5, 3])
                    with col1:
                        if st.button("⬅️ Previous Screen", disabled=(cursor == 0)):
                            st.session_state.orchestrator.handle_ux_previous_screen()
                            st.rerun()

                    with col2:
                        if st.button("Next Screen ➡️", disabled=(cursor >= len(screen_list) - 1)):
                            st.session_state.orchestrator.handle_ux_next_screen()
                            st.rerun()

                elif st.session_state.ux_design_step == 'define_style_guide':
                    st.subheader("Step 5: Define Theming & Style Guide")
                    st.markdown("Describe the desired look and feel of the application (e.g., 'professional, minimalist, dark theme,' 'playful and colorful'). The AI will convert this into a formal style guide.")

                    style_guide_desc = st.text_area(
                        "Describe the desired look and feel:",
                        height=150,
                        key="style_guide_desc"
                    )

                    if st.button("Generate Style Guide"):
                        if style_guide_desc.strip():
                            with st.spinner("Generating Theming & Style Guide..."):
                                st.session_state.orchestrator.handle_style_guide_submission(style_guide_desc)
                            st.rerun()
                        else:
                            st.warning("Please provide a description for the style guide.")

                    # Display the generated style guide for review
                    style_guide_md = st.session_state.orchestrator.active_ux_spec.get('style_guide')
                    if style_guide_md:
                        st.markdown("**Generated Style Guide:**")
                        st.markdown(style_guide_md)
                        st.divider()
                        if st.button("✅ Complete UX/UI Specification", type="primary", use_container_width=True):
                            with st.spinner("Compiling and saving the final UX/UI Specification..."):
                                success = st.session_state.orchestrator.handle_ux_spec_completion()

                            if success:
                                st.success("UX/UI Specification saved. Proceeding to Environment Setup.")
                                # Clean up all session state keys used by the UX/UI phase
                                keys_to_clear = ['ux_design_step', 'personas_text', 'journeys_text', 'screens_text', 'style_guide_desc']
                                for key in keys_to_clear:
                                    if key in st.session_state:
                                        del st.session_state[key]
                                time.sleep(2) # Give user time to read the success message
                                st.rerun()
                            else:
                                error_msg = st.session_state.orchestrator.active_ux_spec.get('error', 'An unknown error occurred.')
                                st.error(f"Failed to finalize the UX/UI Specification: {error_msg}")

        elif current_phase_name == "ENV_SETUP_TARGET_APP":
            st.header(st.session_state.orchestrator.PHASE_DISPLAY_NAMES.get(st.session_state.orchestrator.current_phase))
            st.divider()
            agent = EnvironmentSetupAgent_AppTarget()
            agent.render()

        elif current_phase_name == "SPEC_ELABORATION":
            st.header("Application Specification")
            if 'spec_draft' not in st.session_state:
                st.session_state.spec_draft = None
            if 'spec_step' not in st.session_state:
                st.session_state.spec_step = 'initial_input'
            if 'ai_issues' not in st.session_state:
                st.session_state.ai_issues = None
            if st.session_state.spec_draft is None:
                st.markdown("Please provide the initial specification for your target application.")
                tab1, tab2 = st.tabs(["Upload Specification Documents", "Enter Brief Description"])
                with tab1:
                    uploaded_files = st.file_uploader("Upload Docs", type=["txt", "md", "docx"], accept_multiple_files=True, label_visibility="collapsed")
                    if st.button("Process Uploaded Documents"):
                        if uploaded_files:
                            bootstrap_agent = ProjectBootstrapAgent(db_manager=st.session_state.orchestrator.db_manager)
                            extracted_text, messages, size_error = bootstrap_agent.extract_text_from_files(uploaded_files)
                            for msg in messages:
                                st.warning(msg)
                            if size_error:
                                st.error(size_error)
                            elif extracted_text:
                                st.session_state.spec_draft = extracted_text
                                st.session_state.spec_step = 'complexity_analysis'
                                st.rerun()
                with tab2:
                    brief_desc_input = st.text_area("Brief Description", height=150, key="brief_desc")
                    if st.button("Process Brief Description"):
                        if brief_desc_input:
                            with st.spinner("AI is expanding the description into a draft specification..."):
                                # --- Bugfix: Save the project brief to a file ---
                                project_data_dir = Path(f"data/projects/{st.session_state.orchestrator.project_id}")
                                project_data_dir.mkdir(parents=True, exist_ok=True)
                                brief_file_path = project_data_dir / "project_brief.md"
                                brief_file_path.write_text(brief_desc_input, encoding='utf-8')
                                with st.session_state.orchestrator.db_manager as db:
                                    db.update_project_brief_path(st.session_state.orchestrator.project_id, str(brief_file_path))
                                logging.info(f"Saved text brief to: {brief_file_path}")
                                # --- End of Bugfix ---

                                with st.session_state.orchestrator.db_manager as db:
                                    project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                    is_gui = bool(project_details['is_gui_project']) if project_details else False

                                agent = SpecClarificationAgent(
                                    llm_service=st.session_state.orchestrator.llm_service,
                                    db_manager=st.session_state.orchestrator.db_manager
                                )
                                expanded_text = agent.expand_brief_description(brief_desc_input, is_gui_project=is_gui)
                                st.session_state.spec_draft = expanded_text
                                st.session_state.spec_step = 'complexity_analysis'
                                st.rerun()

            elif st.session_state.spec_step == 'complexity_analysis':
                st.header("Application Specification")
                with st.spinner("Performing complexity and risk analysis on the specification draft..."):
                    try:
                        if not st.session_state.orchestrator.llm_service:
                            st.error("Cannot perform analysis. LLM Service is not configured. Please check Settings.")
                            st.stop()

                        agent = ProjectScopingAgent(llm_service=st.session_state.orchestrator.llm_service)
                        analysis_result = agent.analyze_complexity(st.session_state.spec_draft)

                        if "error" in analysis_result:
                            st.error(f"Failed to analyze project complexity: {analysis_result.get('details', analysis_result['error'])}")
                            st.warning("You can bypass this step and proceed to the manual review.")
                            if st.button("Bypass and Proceed to Manual Review"):
                                st.session_state.spec_step = 'pm_feedback'
                                st.rerun()
                            st.stop()

                        footnote = "\n\n**Note**: This assessment is based on the initial specifications provided to the ASDF. It is a point-in-time analysis and will not be updated if the project requirements evolve."
                        analysis_result_str = json.dumps(analysis_result, indent=4) + footnote

                        with st.session_state.orchestrator.db_manager as db:
                            db.save_complexity_assessment(st.session_state.orchestrator.project_id, analysis_result_str)

                        st.session_state.complexity_analysis = analysis_result
                        st.session_state.spec_step = 'pm_complexity_review'
                        st.rerun()

                    except Exception as e:
                        st.error(f"An unexpected error occurred during complexity analysis: {e}")
                        st.stop()

                        # Append the standard cautionary note
                        footnote = "\n\n**Note**: This assessment is based on the initial specifications provided to the ASDF. It is a point-in-time analysis and will not be updated if the project requirements evolve."
                        analysis_result_str = json.dumps(analysis_result, indent=4) + footnote

                        # Save the full assessment to the database
                        with st.session_state.orchestrator.db_manager as db:
                            db.save_complexity_assessment(st.session_state.orchestrator.project_id, analysis_result_str)

                        # Store the structured result for the UI and transition
                        st.session_state.complexity_analysis = analysis_result
                        st.session_state.spec_step = 'pm_complexity_review'
                        st.rerun()

                    except Exception as e:
                        st.error(f"An unexpected error occurred during complexity analysis: {e}")
                        st.stop()

            elif st.session_state.spec_step == 'pm_complexity_review':
                st.header("Project Complexity & Risk Assessment")
                st.info("The AI has performed a high-level analysis of the specification draft. Please review the assessment below before proceeding.")

                analysis = st.session_state.get('complexity_analysis', {})

                if not analysis:
                    st.error("Complexity analysis result not found. Please go back and try again.")
                    if st.button("Go Back"):
                        st.session_state.spec_step = 'initial_input'
                        if 'spec_draft' in st.session_state:
                            del st.session_state['spec_draft']
                        st.rerun()
                    st.stop()

                # --- Display Risk Assessment ---
                risk = analysis.get('risk_assessment', {})
                st.subheader("Overall Risk Assessment")

                cols = st.columns(2)
                with cols[0]:
                    st.metric("Overall Risk Level", risk.get('overall_risk_level', 'N/A'))
                with cols[1]:
                    st.metric("Token Consumption Outlook", risk.get('token_consumption_outlook', 'N/A'))

                st.write("**Risk Summary:**")
                st.write(risk.get('summary', 'No summary provided.'))

                recommendations = risk.get('recommendations', [])
                if recommendations:
                    st.write("**Recommendations:**")
                    for rec in recommendations:
                        st.warning(f"⚠️ {rec}")

                st.divider()

                # --- Display Complexity Breakdown ---
                with st.expander("Show Detailed Complexity Breakdown"):
                    complexity = analysis.get('complexity_analysis', {})

                    c1, c2 = st.columns(2)
                    with c1:
                        st.subheader("Feature Scope")
                        fs = complexity.get('feature_scope', {})
                        st.metric("Rating", fs.get('rating', 'N/A'))
                        st.caption(fs.get('justification', ''))

                        st.subheader("UI/UX")
                        ui = complexity.get('ui_ux', {})
                        st.metric("Rating", ui.get('rating', 'N/A'))
                        st.caption(ui.get('justification', ''))
                    with c2:
                        st.subheader("Data Schema")
                        ds = complexity.get('data_schema', {})
                        st.metric("Rating", ds.get('rating', 'N/A'))
                        st.caption(ds.get('justification', ''))

                        st.subheader("Integrations")
                        ig = complexity.get('integrations', {})
                        st.metric("Rating", ig.get('rating', 'N/A'))
                        st.caption(ig.get('justification', ''))

                st.divider()

                # --- Final Confirmation ---
                if st.button("Confirm Assessment & Proceed to Review", type="primary", use_container_width=True):
                    del st.session_state.complexity_analysis
                    st.session_state.spec_step = 'pm_feedback'
                    st.rerun()

                st.caption("Note: This assessment is based on the initial specifications provided to the ASDF. It is a point-in-time analysis and will not be updated if the project requirements evolve.")

            elif st.session_state.spec_step == 'pm_feedback':
                st.subheader("Draft Specification - Your Review")
                st.markdown("Please review the initial draft. Provide any corrections, additions, or feedback in the text area below.")
                st.text_area("Current Draft:", value=st.session_state.spec_draft, height=300, key="spec_draft_display", disabled=True)
                pm_feedback_text = st.text_area("Your Feedback and Corrections:", height=200)
                col1, col2, _ = st.columns([1.5, 2.5, 3])
                with col1:
                    if st.button("Submit Feedback & Refine Draft", use_container_width=True):
                        if pm_feedback_text.strip():
                            with st.spinner("AI is refining the draft with your feedback..."):
                                if not st.session_state.orchestrator.llm_service:
                                    st.error("Cannot refine draft: LLM Service is not configured.")
                                    st.stop()

                                with st.session_state.orchestrator.db_manager as db:
                                    project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                    is_gui = bool(project_details['is_gui_project']) if project_details else False

                                agent = SpecClarificationAgent(
                                    llm_service=st.session_state.orchestrator.llm_service,
                                    db_manager=st.session_state.orchestrator.db_manager
                                )
                                refined_draft = agent.refine_specification(
                                    st.session_state.spec_draft,
                                    "PM initial review feedback.",
                                    pm_feedback_text,
                                    is_gui_project=is_gui
                                )
                                st.session_state.spec_draft = refined_draft
                                st.session_state.spec_step = 'pm_review_refined'
                                st.rerun()
                        else:
                            st.warning("Please provide feedback to refine the draft.")
                with col2:
                    if st.button("✅ Approve & Proceed with Review", use_container_width=True, type="primary"):
                        with st.spinner("AI is checking the draft for ambiguities..."):
                            if not st.session_state.orchestrator.llm_service:
                                st.error("Cannot analyze draft: LLM Service is not configured.")
                                st.stop()

                            agent = SpecClarificationAgent(
                                llm_service=st.session_state.orchestrator.llm_service,
                                db_manager=st.session_state.orchestrator.db_manager
                            )
                            issues = agent.identify_potential_issues(st.session_state.spec_draft)
                            st.session_state.ai_issues = issues
                            st.session_state.spec_step = 'ai_clarification'
                            st.rerun()
            elif st.session_state.spec_step == 'pm_review_refined':
                st.subheader("Refined Draft - Your Review")
                st.markdown("The draft has been updated with your feedback. Please review the changes below.")
                st.text_area("Refined Draft:", value=st.session_state.spec_draft, height=300, key="spec_draft_display_refined", disabled=True)
                st.info("You can now either run the AI's ambiguity analysis on this version or provide another round of feedback.")
                col1, col2, _ = st.columns([1.5, 1.5, 3])
                with col1:
                    if st.button("✅ Approve & Proceed to your Review", type="primary", use_container_width=True):
                        with st.spinner("AI is checking the refined draft for ambiguities..."):
                            if not st.session_state.orchestrator.llm_service:
                                st.error("Cannot analyze draft: LLM Service is not configured.")
                                st.stop()

                            agent = SpecClarificationAgent(
                                llm_service=st.session_state.orchestrator.llm_service,
                                db_manager=st.session_state.orchestrator.db_manager
                            )
                            issues = agent.identify_potential_issues(st.session_state.spec_draft)
                            st.session_state.ai_issues = issues
                            st.session_state.spec_step = 'ai_clarification'
                            st.rerun()
                with col2:
                    if st.button("✍️ I Have More Feedback", use_container_width=True):
                        st.session_state.spec_step = 'pm_feedback'
                        st.rerun()
            elif st.session_state.spec_step == 'ai_clarification':
                st.subheader("Specification Refinement - AI Analysis")
                st.text_area("Current Draft:", value=st.session_state.spec_draft, height=300, key="spec_draft_display_2", disabled=True)
                st.divider()
                st.markdown("**AI Analysis Results:**")
                st.info(st.session_state.ai_issues)
                st.markdown("You can now provide clarifications to the AI's points below, or approve the specification if you are satisfied.")
                clarification_text = st.text_area("Your Clarifications:", height=150, key="pm_clarification_text")
                col1, col2, _ = st.columns([1.5, 2, 3])
                with col1:
                    if st.button("Submit Clarifications", use_container_width=True):
                        if clarification_text.strip():
                            with st.spinner("AI is refining the draft and re-analyzing..."):
                                if not st.session_state.orchestrator.llm_service:
                                    st.error("Cannot refine draft: LLM Service is not configured.")
                                    st.stop()

                                with st.session_state.orchestrator.db_manager as db:
                                    project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                    is_gui = bool(project_details['is_gui_project']) if project_details else False

                                agent = SpecClarificationAgent(
                                    llm_service=st.session_state.orchestrator.llm_service,
                                    db_manager=st.session_state.orchestrator.db_manager
                                )
                                st.session_state.orchestrator.capture_spec_clarification_learning(
                                    problem_context=st.session_state.ai_issues,
                                    solution_text=clarification_text,
                                    spec_text=st.session_state.spec_draft
                                )
                                refined_draft = agent.refine_specification(
                                    st.session_state.spec_draft,
                                    st.session_state.ai_issues,
                                    clarification_text,
                                    is_gui_project=is_gui
                                )
                                st.session_state.spec_draft = refined_draft
                                issues = agent.identify_potential_issues(st.session_state.spec_draft)
                                st.session_state.ai_issues = issues
                                st.rerun()
                        else:
                            st.warning("Please enter your clarifications before submitting.")
                with col2:
                    if st.button("✅ Approve Specification and Proceed", type="primary", use_container_width=True):
                        with st.spinner("Finalizing and saving specification..."):
                            st.session_state.orchestrator.finalize_and_save_app_spec(st.session_state.spec_draft)
                            # Clean up UI state
                            keys_to_clear = ['spec_draft', 'spec_step', 'ai_issues', 'brief_desc', 'pm_clarification_text']
                            for key in keys_to_clear:
                                if key in st.session_state:
                                    del st.session_state[key]
                            st.rerun()

        elif current_phase_name == "TECHNICAL_SPECIFICATION":
            st.header(st.session_state.orchestrator.PHASE_DISPLAY_NAMES.get(st.session_state.orchestrator.current_phase))
            if 'tech_spec_step' not in st.session_state:
                st.session_state.tech_spec_step = 'initial_choice'
            if 'tech_spec_draft' not in st.session_state:
                st.session_state.tech_spec_draft = ""
            if 'target_os' not in st.session_state:
                st.session_state.target_os = "Linux"

            if st.session_state.tech_spec_step == 'initial_choice':
                st.markdown("First, select the target Operating System for the application.")
                st.session_state.target_os = st.selectbox(
                    "Select Target Operating System:",
                    ["Linux", "Windows", "macOS"],
                    index=["Linux", "Windows", "macOS"].index(st.session_state.target_os)
                )
                st.divider()
                st.markdown("Next, choose how you would like to create the Technical Specification document.")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🤖 Let ASDF Propose a Tech Stack", use_container_width=True, type="primary"):
                        with st.spinner("AI is analyzing the specification and generating a proposal..."):
                            try:
                                if not st.session_state.orchestrator.llm_service:
                                    st.error("Cannot generate proposal: LLM Service is not configured.")
                                else:
                                    with st.session_state.orchestrator.db_manager as db:
                                        project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                        final_spec_text = project_details['final_spec_text']

                                    agent = TechStackProposalAgent(llm_service=st.session_state.orchestrator.llm_service)
                                    proposal = agent.propose_stack(final_spec_text, st.session_state.target_os)
                                    st.session_state.tech_spec_draft = proposal
                                    st.session_state.tech_spec_step = 'pm_review'
                                    st.rerun()
                            except Exception as e:
                                st.error(f"Failed to generate proposal: {e}")
                with col2:
                    if st.button("✍️ Set Technology Choices/Guidelines", use_container_width=True):
                        st.session_state.tech_spec_step = 'pm_define'
                        st.rerun()

            elif st.session_state.tech_spec_step == 'pm_define':
                st.subheader("Define Technical Specification")
                st.markdown("Please provide your key technology choices or guidelines below (e.g., 'Use Python with a Flask backend and a SQLite database'). The ASDF will use your input to generate the full technical specification document.")
                st.session_state.tech_spec_draft = st.text_area(
                    "Your Technology Guidelines:",
                    value=st.session_state.tech_spec_draft,
                    height=200
                )
                if st.button("Generate Full Specification from My Input", type="primary"):
                    if st.session_state.tech_spec_draft.strip():
                        with st.spinner("AI is expanding your input into a full technical specification..."):
                            try:
                                if not st.session_state.orchestrator.llm_service:
                                    st.error("Cannot generate specification: LLM Service is not configured.")
                                else:
                                    with st.session_state.orchestrator.db_manager as db:
                                        project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                        context = (
                                            f"{project_details['final_spec_text']}\n\n"
                                            f"--- PM Directive for Technology Stack ---\n{st.session_state.tech_spec_draft}"
                                        )
                                    agent = TechStackProposalAgent(llm_service=st.session_state.orchestrator.llm_service)
                                    proposal = agent.propose_stack(context, st.session_state.target_os)
                                    st.session_state.tech_spec_draft = proposal
                                    st.session_state.tech_spec_step = 'pm_review'
                                    st.rerun()
                            except Exception as e:
                                st.error(f"Failed to generate specification from your input: {e}")
                    else:
                        st.warning("Please provide your technology guidelines before generating the specification.")

            elif st.session_state.tech_spec_step == 'pm_review':
                st.subheader("Review Technical Specification")
                st.markdown("Please review the draft below. You can either approve it or provide feedback for refinement.")
                st.session_state.tech_spec_draft = st.text_area(
                    "Technical Specification Draft",
                    value=st.session_state.tech_spec_draft,
                    height=400,
                    key="tech_spec_draft_display",
                    disabled=False
                )
                feedback_text = st.text_area("Your Feedback and Refinements:", height=150)
                col1, col2, _ = st.columns([1.5, 2, 3])
                with col1:
                    if st.button("Submit Feedback & Refine", use_container_width=True):
                        if feedback_text.strip():
                            with st.spinner("AI is refining the proposal based on your feedback..."):
                                try:
                                    if not st.session_state.orchestrator.llm_service:
                                        st.error("Cannot refine proposal: LLM Service is not configured.")
                                    else:
                                        with st.session_state.orchestrator.db_manager as db:
                                            project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                            context = (
                                                f"{project_details['final_spec_text']}\n\n"
                                                f"--- Current Draft to Refine ---\n{st.session_state.tech_spec_draft}\n\n"
                                                f"--- PM Feedback for Refinement ---\n{feedback_text}"
                                            )
                                        agent = TechStackProposalAgent(llm_service=st.session_state.orchestrator.llm_service)
                                        proposal = agent.propose_stack(context, st.session_state.target_os)
                                        st.session_state.tech_spec_draft = proposal
                                        st.rerun()
                                except Exception as e:
                                    st.error(f"Failed to generate proposal: {e}")
                        else:
                            st.warning("Please enter feedback before submitting for refinement.")
                with col2:
                    is_disabled = not st.session_state.tech_spec_draft.strip()
                    if st.button("✅ Approve Specification", use_container_width=True, type="primary", disabled=is_disabled):
                        with st.spinner("Saving technical specification..."):
                            st.session_state.orchestrator.finalize_and_save_tech_spec(
                                st.session_state.tech_spec_draft,
                                st.session_state.target_os
                            )
                        keys_to_clear = ['tech_spec_draft', 'target_os', 'tech_spec_step']
                        for key in keys_to_clear:
                            if key in st.session_state:
                                del st.session_state[key]
                        st.rerun()

        elif current_phase_name == "BUILD_SCRIPT_SETUP":
            st.header(st.session_state.orchestrator.PHASE_DISPLAY_NAMES.get(st.session_state.orchestrator.current_phase))
            st.markdown("The technical specification is complete. Now, let's establish the build script for the project.")
            st.info("This script (e.g., `requirements.txt`, `pom.xml`) manages project dependencies and how the application is built.")

            with st.session_state.orchestrator.db_manager as db:
                project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                tech_spec_text = project_details['tech_spec_text']
                target_os = project_details['target_os']

            if not target_os:
                st.error("Cannot proceed: Target OS has not been set in the previous step.")
            else:
                st.write(f"Generating options for Target OS: **{target_os}**")
                st.divider()
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🤖 Auto-Generate Build Script", use_container_width=True, type="primary"):
                        with st.spinner(f"Generating standard build script for the specified tech stack..."):
                            try:
                                if not st.session_state.orchestrator.llm_service:
                                    st.error("Cannot generate script: LLM Service is not configured.")
                                else:
                                    from agents.agent_build_script_generator import BuildScriptGeneratorAgent
                                    agent = BuildScriptGeneratorAgent(llm_service=st.session_state.orchestrator.llm_service)
                                    script_info = agent.generate_script(tech_spec_text, target_os)
                                    if script_info:
                                        filename, content = script_info
                                        project_root = Path(project_details['project_root_folder'])
                                        (project_root / filename).write_text(content, encoding='utf-8')
                                        with st.session_state.orchestrator.db_manager as db:
                                            db.update_project_build_automation_status(st.session_state.orchestrator.project_id, True)
                                        st.success(f"Generated and saved `{filename}` to the project root.")
                                        time.sleep(2)
                                        st.session_state.orchestrator.set_phase("TEST_ENVIRONMENT_SETUP")
                                        st.rerun()
                                    else:
                                        st.error(f"The AI was unable to generate a build script. Please proceed manually.")
                            except Exception as e:
                                st.error(f"An error occurred during build script generation: {e}")
                with col2:
                    if st.button("✍️ I Will Create it Manually", use_container_width=True):
                        with st.session_state.orchestrator.db_manager as db:
                            db.update_project_build_automation_status(st.session_state.orchestrator.project_id, False)
                        st.info("Acknowledged. You will be responsible for the project's build script.")
                        time.sleep(2)
                        st.session_state.orchestrator.set_phase("TEST_ENVIRONMENT_SETUP")
                        st.rerun()

        elif current_phase_name == "TEST_ENVIRONMENT_SETUP":
            st.header(st.session_state.orchestrator.PHASE_DISPLAY_NAMES.get(st.session_state.orchestrator.current_phase))
            st.markdown("Please follow the steps below to set up the necessary testing frameworks for your project's technology stack.")

            if 'setup_tasks' not in st.session_state:
                with st.spinner("Analyzing technical spec to generate setup steps..."):
                    st.session_state.setup_tasks = st.session_state.orchestrator.start_test_environment_setup()
                    st.session_state.current_setup_step = 0
                    st.session_state.setup_help_text = None

            tasks = st.session_state.setup_tasks
            if tasks is None:
                st.error("Could not generate setup tasks. Please check the logs and ensure the technical specification is complete.")
            elif st.session_state.current_setup_step >= len(tasks):
                st.success("All setup steps have been actioned.")
                st.markdown("Please confirm the final command that should be used to run all automated tests for this project.")

                with st.session_state.orchestrator.db_manager as db:
                     project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                     tech_spec_text = project_details['tech_spec_text']

                if 'suggested_test_command' not in st.session_state:
                    if st.session_state.orchestrator.llm_service and tech_spec_text:
                        agent = VerificationAgent_AppTarget(llm_service=st.session_state.orchestrator.llm_service)
                        details = agent._get_test_execution_details(tech_spec_text)
                        st.session_state.suggested_test_command = details.get("command") if details else "pytest"
                    else:
                        st.session_state.suggested_test_command = "pytest"

                confirmed_command = st.text_input(
                    "Test Execution Command:",
                    value=st.session_state.suggested_test_command
                )
                if st.button("Finalize Test Environment Setup", type="primary"):
                    if confirmed_command.strip():
                        if st.session_state.orchestrator.finalize_test_environment_setup(confirmed_command):
                            keys_to_clear = ['setup_tasks', 'current_setup_step', 'setup_help_text', 'suggested_test_command']
                            for key in keys_to_clear:
                                if key in st.session_state:
                                    del st.session_state[key]
                            st.rerun()
                        else:
                            st.error("Failed to finalize setup. Please check the logs.")
                    else:
                        st.warning("The test execution command cannot be empty.")
            else:
                current_step_index = st.session_state.current_setup_step
                task = tasks[current_step_index]
                st.subheader(f"Step {current_step_index + 1} of {len(tasks)}: {task.get('tool_name', 'Unnamed Step')}")
                instructions_text = task.get('instructions', 'No instructions provided.')
                # Sanitize LLM output to prevent inconsistent font sizes by removing markdown headings
                sanitized_instructions = "\n".join([line.lstrip('# ') for line in instructions_text.split('\n')])
                st.markdown(sanitized_instructions)
                if st.session_state.setup_help_text:
                    with st.chat_message("assistant", avatar="❓"):
                        st.info(st.session_state.setup_help_text)
                st.divider()
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("✅ Done, Next Step", use_container_width=True, type="primary"):
                        st.session_state.current_setup_step += 1
                        st.session_state.setup_help_text = None
                        st.rerun()
                with col2:
                    if st.button("❓ I Need Help", use_container_width=True):
                        with st.spinner("Getting more details..."):
                            st.session_state.setup_help_text = st.session_state.orchestrator.get_help_for_setup_task(task.get('instructions', ''))
                        st.rerun()
                with col3:
                    if st.button("⚠️ Ignore & Continue", use_container_width=True):
                        task = st.session_state.setup_tasks[st.session_state.current_setup_step]
                        st.session_state.orchestrator.handle_ignore_setup_task(task)
                        st.session_state.current_setup_step += 1
                        st.session_state.setup_help_text = None
                        st.rerun()

        elif current_phase_name == "CODING_STANDARD_GENERATION":
            st.header("Coding Standard Generation")
            if 'coding_standard_step' not in st.session_state:
                st.session_state.coding_standard_step = 'initial'
            if 'coding_standard_draft' not in st.session_state:
                st.session_state.coding_standard_draft = ""

            if st.session_state.coding_standard_step == 'initial':
                st.markdown("Generate a project-specific coding standard based on the technical specification. This standard will be enforced by all code-generating agents.")
                if st.button("Generate Coding Standard Draft", type="primary"):
                    with st.spinner("AI is generating the coding standard..."):
                        try:
                            if not st.session_state.orchestrator.llm_service:
                                st.error("Cannot generate standard: LLM Service is not configured.")
                            else:
                                with st.session_state.orchestrator.db_manager as db:
                                    project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                    tech_spec = project_details['tech_spec_text']

                                if not tech_spec:
                                    st.error("Cannot generate standard: Technical Specification is missing.")
                                else:
                                    from agents.agent_coding_standard_app_target import CodingStandardAgent_AppTarget
                                    agent = CodingStandardAgent_AppTarget(llm_service=st.session_state.orchestrator.llm_service)
                                    standard = agent.generate_standard(tech_spec)
                                    st.session_state.coding_standard_draft = standard
                                    st.session_state.coding_standard_step = 'pm_review'
                                    st.rerun()
                        except Exception as e:
                            st.error(f"Failed to generate coding standard: {e}")

            elif st.session_state.coding_standard_step == 'pm_review':
                st.subheader("Review Coding Standard")
                st.markdown("Please review the draft below. You can either approve it or provide feedback for refinement.")
                st.session_state.coding_standard_draft = st.text_area(
                    "Coding Standard Draft",
                    value=st.session_state.coding_standard_draft,
                    height=400,
                    key="coding_standard_display"
                )
                feedback_text = st.text_area("Your Feedback and Refinements:", height=150)

                col1, col2, _ = st.columns([1.5, 2, 3])
                with col1:
                    if st.button("Submit Feedback & Refine", use_container_width=True):
                        if feedback_text.strip():
                            st.warning("Feedback refinement for the Coding Standard is not yet implemented.")
                        else:
                            st.warning("Please enter feedback before submitting for refinement.")
                with col2:
                    is_disabled = not st.session_state.coding_standard_draft.strip()
                    if st.button("✅ Approve Coding Standard", use_container_width=True, type="primary", disabled=is_disabled):
                        with st.spinner("Saving coding standard..."):
                            st.session_state.orchestrator.finalize_and_save_coding_standard(
                                st.session_state.coding_standard_draft
                            )
                        keys_to_clear = ['coding_standard_draft', 'coding_standard_step']
                        for key in keys_to_clear:
                            if key in st.session_state:
                                del st.session_state[key]
                        st.rerun()

        elif current_phase_name == "PLANNING":
            st.header("Strategic Development Planning")
            if 'development_plan' not in st.session_state:
                st.session_state.development_plan = None
            if st.session_state.development_plan:
                st.subheader("Generated Development Plan")
                st.markdown("Please review the generated plan. To proceed, approve it below. To generate a new version, click the 'Re-generate' button.")
                st.json(st.session_state.development_plan)
                st.divider()
                col1, col2, col3 = st.columns([1.5, 1, 1.5])
                with col1:
                    if st.button("✅ Approve Plan & Proceed to Development", type="primary"):
                        with st.spinner("Saving plan and transitioning to development phase..."):
                            success, message = st.session_state.orchestrator.finalize_and_save_dev_plan(
                                st.session_state.development_plan
                            )

                        if success:
                            st.toast(message)
                            # Clean up UI state
                            if 'development_plan' in st.session_state:
                                del st.session_state['development_plan']
                            st.rerun()
                        else:
                            st.error(message)
                with col2:
                    report_generator = ReportGeneratorAgent()
                    dev_plan_docx_bytes = report_generator.generate_text_document_docx(
                        title=f"Sequential Development Plan - {st.session_state.orchestrator.project_name}",
                        content=st.session_state.development_plan,
                        is_code=True
                    )
                    st.download_button(
                        label="📄 Print to .docx",
                        data=dev_plan_docx_bytes,
                        file_name=f"Development_Plan_{st.session_state.orchestrator.project_id}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with col3:
                    if st.button("🔄 Re-generate Development Plan"):
                         st.session_state.development_plan = None
                         st.rerun()
            else:
                st.info("Click the button below to generate a detailed, sequential development plan based on the finalized specifications.")
                if st.button("Generate Development Plan", type="primary"):
                    with st.spinner("AI is generating the development plan... This may take a few moments."):
                        try:
                            if not st.session_state.orchestrator.llm_service:
                                st.error("Could not generate plan: LLM Service is not configured. Please check Settings.")
                                st.stop()

                            with st.session_state.orchestrator.db_manager as db:
                                project_details = db.get_project_by_id(st.session_state.orchestrator.project_id)
                                final_spec = project_details['final_spec_text']
                                tech_spec = project_details['tech_spec_text']

                            if not all([final_spec, tech_spec]):
                                st.error("Could not generate plan: Missing Final Specification or Technical Specification in the database.")
                            else:
                                agent = PlanningAgent_AppTarget(llm_service=st.session_state.orchestrator.llm_service, db_manager=st.session_state.orchestrator.db_manager)
                                response_json_str = agent.generate_development_plan(final_spec, tech_spec)
                                response_data = json.loads(response_json_str)
                                if "error" in response_data:
                                    st.error(f"Failed to generate plan: {response_data['error']}")
                                else:
                                    plan_data = response_data.get("development_plan")
                                    main_executable = response_data.get("main_executable_file")
                                    if not plan_data or not main_executable:
                                        st.error("Failed to generate a valid plan and executable name from the AI.")
                                    else:
                                        with st.session_state.orchestrator.db_manager as db:
                                            db.update_project_apex_file(st.session_state.orchestrator.project_id, main_executable)
                                        st.session_state.development_plan = json.dumps(response_data, indent=4)
                                        st.rerun()
                        except Exception as e:
                            st.error(f"An unexpected error occurred: {e}")

        elif current_phase_name == "GENESIS":
            st.header("Iterative Component Development")
            if not st.session_state.orchestrator.active_plan:
                st.warning("No active development plan is loaded.")
                st.info("Please go to the 'Planning' phase to generate a development plan.")
                if st.button("⬅️ Go to Planning Phase"):
                    st.session_state.orchestrator.set_phase("PLANNING")
                    st.rerun()
            else:
                st.subheader("PM Checkpoint")
                task = st.session_state.orchestrator.get_current_task_details()
                total_tasks = len(st.session_state.orchestrator.active_plan)
                cursor = st.session_state.orchestrator.active_plan_cursor
                if task:
                    st.progress(cursor / total_tasks, text=f"Executing Task {cursor + 1} of {total_tasks}")
                    st.info(f"""
                    Next component in the plan is: **'{task.get('component_name')}'** (based on micro-spec `{task.get('micro_spec_id')}`).
                    How would you like to proceed?
                    """)
                else:
                    st.success(f"All {total_tasks} development tasks are complete.")
                    st.info("Click 'Proceed' to finalize development and begin the Integration & Validation phase.")
                # Use a session state flag to disable buttons after a click
                if 'dev_step_in_progress' not in st.session_state:
                    st.session_state.dev_step_in_progress = False

                def set_step_in_progress():
                    st.session_state.dev_step_in_progress = True

                col1, col2, _ = st.columns([1.5, 1.5, 3])
                with col1:
                    if st.button(
                        "▶️ Proceed",
                        use_container_width=True,
                        type="primary",
                        on_click=set_step_in_progress,
                        disabled=st.session_state.dev_step_in_progress
                    ):
                        with st.status("Executing next development step...", expanded=True) as status:
                            st.session_state.orchestrator.handle_proceed_action(status_ui_object=status)
                            status.update(label="Step completed successfully!", state="complete", expanded=False)
                        del st.session_state.dev_step_in_progress # Reset flag
                        st.rerun()
                with col2:
                    st.button(
                        "⏹️ Stop & Export",
                        use_container_width=True,
                        on_click=set_step_in_progress,
                        disabled=st.session_state.dev_step_in_progress
                    )
                    if st.session_state.dev_step_in_progress:
                        st.session_state.show_export_confirmation = True
                        st.rerun()

        elif current_phase_name == "AWAITING_INTEGRATION_CONFIRMATION":
            st.header("Integration Warning")
            st.warning(
                "**The development phase is complete, but one or more components have known issues or failed their unit tests.**"
            )
            st.markdown(
                "Proceeding with integration may result in an unstable or non-functional build. It is recommended to fix these issues before continuing."
            )
            issues = st.session_state.orchestrator.task_awaiting_approval.get("known_issues", [])
            if issues:
                st.subheader("Components with Known Issues:")
                for issue in issues:
                    st.error(f"- **{issue.get('artifact_name')}** (Status: {issue.get('status')})")
            st.divider()
            st.markdown("Do you want to proceed with integration anyway?")
            col1, col2, _ = st.columns([1.5, 2, 3])
            with col1:
                if st.button("✅ Yes, Proceed to Integration", type="primary"):
                    st.session_state.orchestrator.task_awaiting_approval = None
                    st.session_state.orchestrator._run_integration_and_ui_testing_phase()
                    st.rerun()
            with col2:
                if st.button("❌ No, Return to Development"):
                    st.session_state.orchestrator.task_awaiting_approval = None
                    st.session_state.orchestrator.set_phase("GENESIS")
                    st.rerun()

        elif current_phase_name == "MANUAL_UI_TESTING":
            st.header("Manual UI Testing")
            st.info(
                "The automated integration and build were successful. "
                "The UI Test Plan has been generated based on the project specification."
            )
            st.markdown(
                "**Your action is required:**\n\n"
                "1.  Navigate to the **Documents** page from the sidebar.\n"
                "2.  Download the **UI Test Plan** document.\n"
                "3.  Execute the tests as described and fill in the results.\n"
                "4.  Upload the completed document below to trigger the evaluation."
            )
            st.divider()
            uploaded_file = st.file_uploader(
                "Upload Completed UI Test Plan Results",
                type=['txt', 'md', 'docx']
            )
            if uploaded_file is not None:
                if st.button("Process Test Results", type="primary"):
                    with st.spinner("Reading and evaluating test results..."):
                        content = ""
                        try:
                            if uploaded_file.name.endswith('.docx'):
                                import docx
                                doc = docx.Document(uploaded_file)
                                content = "\n".join([p.text for p in doc.paragraphs])
                            else: # For .txt and .md files
                                content = uploaded_file.getvalue().decode("utf-8")
                            st.session_state.orchestrator.handle_ui_test_result_upload(content)
                            st.success("Test results submitted for evaluation. The system will now process any failures.")
                            time.sleep(2)
                            st.rerun()
                        except Exception as e:
                            st.error(f"Failed to process file: {e}")
            st.divider()
            st.markdown("If all tests have passed and been fixed, you can proceed.")
            if st.button("Complete Project and Return to Idle"):
                st.session_state.orchestrator.set_phase("IDLE")
                st.toast("Project phase complete!")
                st.rerun()

        elif current_phase_name == "AWAITING_INTEGRATION_RESOLUTION":
            st.header("Integration & Verification Failed")
            st.error("The automated integration and verification process could not be completed due to a system-level or environment error.")
            if st.session_state.orchestrator.task_awaiting_approval:
                failure_reason = st.session_state.orchestrator.task_awaiting_approval.get("failure_reason", "No specific reason was provided.")
                st.markdown("**Failure Details:**")
                st.code(failure_reason, language='text')
            st.warning("This is often caused by a missing testing framework (like 'pytest') in your environment or a configuration issue. You cannot re-run the automated test, but you can acknowledge the failure and proceed to manual UI testing.")
            if st.button("Acknowledge Failure & Proceed to Manual UI Testing", type="primary", use_container_width=True):
                st.session_state.orchestrator.handle_acknowledge_integration_failure()
                st.rerun()

        elif current_phase_name == "AWAITING_PM_TRIAGE_INPUT":
            st.header("Interactive Debugging Triage")
            st.warning(
                "**Action Required:** The automated triage system could not determine the "
                "root cause of the last failure from the available logs."
            )
            st.markdown(
                "Please describe the failure in as much detail as possible below. "
                "This description will be used to generate a fix."
            )
            st.divider()
            if 'pm_triage_input' not in st.session_state:
                st.session_state.pm_triage_input = ""
            pm_error_description = st.text_area(
                "Manual Error Description:",
                height=200,
                key="pm_triage_input"
            )
            col1, col2, _ = st.columns([1.5, 2, 3])
            with col1:
                if st.button("Submit Description for Analysis", type="primary"):
                    if pm_error_description.strip():
                        with st.spinner("Submitting your description for analysis..."):
                            st.session_state.orchestrator.handle_pm_triage_input(pm_error_description)
                        st.toast("Description submitted. The factory will now attempt to plan a fix.")
                        time.sleep(2)
                        st.rerun()
                    else:
                        st.error("Please provide a description of the error before submitting.")
            with col2:
                if st.button("Cancel and Return to Checkpoint"):
                    st.session_state.orchestrator.set_phase("GENESIS")
                    st.toast("Returning to the main development checkpoint.")
                    st.rerun()

        elif current_phase_name == "DEBUG_PM_ESCALATION":
            st.header("Debug Escalation")
            st.error(
                "**Action Required:** The factory has been unable to automatically fix a persistent bug after multiple attempts."
            )
            if st.session_state.orchestrator.task_awaiting_approval:
                failure_context = st.session_state.orchestrator.task_awaiting_approval.get("failure_log", "No specific failure log was captured.")
                with st.expander("Click to view failure details"):
                    st.code(failure_context, language='text')
            st.markdown(
                "Please choose how you would like to proceed. Your selection will determine the next steps for the factory."
            )
            st.divider()
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("🔄 Retry Automated Fix", use_container_width=True, help="This will reset the counter and run the entire triage and fix pipeline again from a clean state."):
                    st.session_state.orchestrator.handle_pm_debug_choice("RETRY")
                    st.rerun()
            with col2:
                if st.button("⏸️ Pause for Manual Fix", use_container_width=True, help="This will pause the factory, allowing you to manually investigate and fix the code in your own editor."):
                    st.session_state.orchestrator.handle_pm_debug_choice("MANUAL_PAUSE")
                    st.success("Project paused. You can now manually edit the code. The factory will remain idle.")
                    st.rerun()
            with col3:
                if st.button("🚫 Ignore Bug & Proceed", use_container_width=True, help="Acknowledge the bug but proceed with the next task in the development plan. The bug will be noted but not fixed now."):
                    st.session_state.orchestrator.handle_pm_debug_choice("IGNORE")
                    st.toast("Bug ignored. Proceeding with the next development task.")
                    st.rerun()

        elif current_phase_name == "INTEGRATION_AND_VERIFICATION":
            st.header("Automated Integration & Verification")
            st.info("The factory is now integrating all newly developed components, performing a full system build, and running verification tests.")
            with st.spinner("Running automated integration... This may take a moment."):
                st.session_state.orchestrator._run_integration_and_verification_phase()
            st.success("Integration and verification process complete.")
            time.sleep(2) # Pause for 2 seconds to allow the user to read the message
            st.rerun()

        elif current_phase_name == "AWAITING_PM_DECLARATIVE_CHECKPOINT":
            st.header("PM Checkpoint: High-Risk Change Detected")
            st.warning("The development plan requires a modification to a declarative file (e.g., build script, database schema, config file). Please review the proposed change below.")
            task = st.session_state.orchestrator.task_awaiting_approval
            if task:
                st.markdown(f"**File to Modify:** `{task.get('component_file_path')}`")
                st.markdown(f"**Component:** `{task.get('component_name')}`")
                st.subheader("Proposed Change Snippet")
                st.code(task.get('task_description'), language='diff')
                st.divider()
                st.markdown("How would you like to proceed with this change?")
                col1, col2, _ = st.columns([1, 1, 3])
                with col1:
                    if st.button("✅ Execute Automatically", use_container_width=True, type="primary"):
                        st.session_state.orchestrator.handle_declarative_checkpoint_decision("EXECUTE_AUTOMATICALLY")
                        st.toast(f"Executing change for {task.get('component_name')}...")
                        st.rerun()
                with col2:
                    if st.button("✍️ I will Execute Manually", use_container_width=True):
                        st.session_state.orchestrator.handle_declarative_checkpoint_decision("WILL_EXECUTE_MANUALLY")
                        st.toast("Acknowledged. Please apply the change manually.")
                        st.rerun()
            else:
                st.error("Could not retrieve the task awaiting approval. Returning to Genesis phase.")
                st.session_state.orchestrator.set_phase("GENESIS")
                if st.button("Go Back"):
                    st.rerun()

        elif current_phase_name == "RAISING_CHANGE_REQUEST":
            st.header("Raise New Change Request")
            st.info("First, select the type of change you are requesting.")
            if 'cr_type' not in st.session_state:
                st.session_state.cr_type = "Functional Enhancement"
            if 'cr_description' not in st.session_state:
                st.session_state.cr_description = ""
            if 'spec_correction_text' not in st.session_state:
                st.session_state.spec_correction_text = ""
            st.session_state.cr_type = st.radio(
                "Select Change Request Type:",
                ["Functional Enhancement", "Specification Correction"],
                horizontal=True,
                key="cr_type_radio"
            )
            st.divider()
            if st.session_state.cr_type == "Functional Enhancement":
                st.markdown("Please provide a detailed description of the new feature or change in functionality.")
                st.session_state.cr_description = st.text_area(
                    "Change Request Description:",
                    value=st.session_state.cr_description,
                    height=250
                )
                if st.button("Save Change Request", type="primary"):
                    if st.session_state.cr_description.strip():
                        st.session_state.orchestrator.save_new_change_request(st.session_state.cr_description, "CHANGE_REQUEST")
                        st.rerun()
                    else:
                        st.warning("The description cannot be empty.")
            elif st.session_state.cr_type == "Specification Correction":
                st.markdown("Paste the **full, corrected text** of the specification below. The system will analyze the differences to create a linked implementation plan.")
                if not st.session_state.spec_correction_text:
                     with st.session_state.orchestrator.db_manager as db:
                        project_docs = db.get_project_by_id(st.session_state.orchestrator.project_id)
                        st.session_state.spec_correction_text = project_docs['final_spec_text'] if project_docs else ""
                st.session_state.spec_correction_text = st.text_area(
                    "Full Corrected Specification Text:",
                    value=st.session_state.spec_correction_text,
                    height=400
                )
                if st.button("Save Specification Change", type="primary"):
                    if st.session_state.spec_correction_text.strip():
                        st.session_state.orchestrator.save_spec_correction_cr(st.session_state.spec_correction_text)
                        st.rerun()
                    else:
                        st.warning("The specification text cannot be empty.")
            st.divider()
            if st.button("Cancel"):
                st.session_state.orchestrator.set_phase("GENESIS")
                st.rerun()

        elif current_phase_name == "AWAITING_INITIAL_IMPACT_ANALYSIS":
            st.header("New Change Request Logged")
            st.success("The new Change Request has been saved to the register.")
            st.markdown("Would you like to perform a high-level impact analysis on this new CR now?")
            try:
                latest_cr_id = st.session_state.orchestrator.get_all_change_requests()[0]['cr_id']
            except (IndexError, TypeError):
                st.error("Could not retrieve the newly created Change Request. Returning to checkpoint.")
                if st.button("Back to Checkpoint"):
                    st.session_state.orchestrator.set_phase("GENESIS")
                    st.rerun()
                st.stop()
            col1, col2, _ = st.columns([1, 1, 5])
            with col1:
                if st.button("Yes, Run Analysis", type="primary", use_container_width=True):
                    with st.spinner(f"Running impact analysis for CR-{latest_cr_id}..."):
                        st.session_state.orchestrator.handle_run_impact_analysis_action(latest_cr_id)
                    st.toast(f"Impact analysis complete for CR-{latest_cr_id}.")
                    st.session_state.orchestrator.set_phase("GENESIS")
                    st.rerun()
            with col2:
                if st.button("No, Later", use_container_width=True):
                    st.toast("Acknowledged. You can run the analysis later from the CR Register.")
                    st.session_state.orchestrator.set_phase("GENESIS")
                    st.rerun()

        elif current_phase_name == "IMPLEMENTING_CHANGE_REQUEST":
            st.header("Implement Requested Change")
            st.markdown("Select a Change Request from the register below to view available actions.")
            change_requests = st.session_state.orchestrator.get_all_change_requests()
            if not change_requests:
                st.warning("There are no change requests in the register for this project.")
            else:
                cr_data_for_df = []
                for cr in change_requests:
                    cr_data_for_df.append({
                        "ID": cr['cr_id'],
                        "Type": cr['request_type'].replace('_', ' ').title(),
                        "Status": cr['status'],
                        "Severity/Impact": cr['impact_rating'],
                        "Description": cr['description'],
                        "Analysis Summary": cr['impact_analysis_details']
                    })
                df = pd.DataFrame(cr_data_for_df)
                column_order = ["ID", "Type", "Status", "Severity/Impact", "Description", "Analysis Summary"]
                st.dataframe(df[column_order], use_container_width=True, hide_index=True)
                cr_ids = [cr['cr_id'] for cr in change_requests]
                selected_cr_id_str = st.selectbox("Select a Change Request ID to action:", options=[""] + [str(i) for i in cr_ids])
                if selected_cr_id_str:
                    selected_cr_id = int(selected_cr_id_str)
                    selected_cr = next((cr for cr in change_requests if cr['cr_id'] == selected_cr_id), None)
                    st.subheader(f"Actions for CR-{selected_cr_id}")
                    is_raised_status = selected_cr['status'] == 'RAISED'
                    is_impact_analyzed = selected_cr['status'] == 'IMPACT_ANALYZED'
                    genesis_is_complete = st.session_state.orchestrator.is_genesis_complete
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        if st.button("✏️ Edit CR", use_container_width=True, disabled=not is_raised_status, help="You can only edit a CR before its impact has been analyzed."):
                            st.session_state.orchestrator.handle_edit_cr_action(selected_cr_id)
                            st.rerun()
                    with col2:
                        if st.button("🗑️ Delete CR", use_container_width=True, disabled=not is_raised_status, help="You can only delete a CR before its impact has been analyzed."):
                            with st.popover("Confirm Deletion"):
                                st.write(f"Are you sure you want to permanently delete CR-{selected_cr_id}?")
                                if st.button("Yes, Confirm Delete", type="primary"):
                                    st.session_state.orchestrator.handle_delete_cr_action(selected_cr_id)
                                    st.toast(f"Change Request {selected_cr_id} deleted.")
                                    st.rerun()
                    with col3:
                        if st.button("🔬 Run Impact Analysis", use_container_width=True, disabled=not is_raised_status, help="Run analysis to determine the scope and impact of the change."):
                            with st.spinner(f"Running impact analysis for CR-{selected_cr_id}..."):
                                st.session_state.orchestrator.handle_run_impact_analysis_action(selected_cr_id)
                            st.toast(f"Impact analysis complete for CR-{selected_cr_id}.")
                            st.rerun()
                    with col4:
                        if st.button("▶️ Implement CR", use_container_width=True, type="primary", disabled=not (is_impact_analyzed and genesis_is_complete), help="Implementation is enabled only after impact analysis is run and main development is complete."):
                            with st.spinner(f"Generating refactoring plan for CR-{selected_cr_id}..."):
                                st.session_state.orchestrator.handle_implement_cr_action(selected_cr_id)
                            st.rerun()
                if not st.session_state.orchestrator.is_genesis_complete:
                    st.info("Note: CR implementation is enabled after the main development plan is fully completed.")
            st.divider()
            if st.button("⬅️ Back to Main Checkpoint"):
                st.session_state.orchestrator.set_phase("GENESIS")
                st.rerun()

        elif current_phase_name == "EDITING_CHANGE_REQUEST":
            st.header("Edit Change Request")
            cr_details = st.session_state.orchestrator.get_active_cr_details_for_edit()
            if not cr_details:
                st.error("Error: Could not load Change Request details for editing.")
                if st.button("⬅️ Back to Register"):
                    st.session_state.orchestrator.cancel_cr_edit()
                    st.rerun()
            else:
                if 'cr_edit_description' not in st.session_state:
                    st.session_state.cr_edit_description = cr_details.get('description', '')
                st.markdown(f"You are editing **CR-{cr_details['cr_id']}**.")
                st.session_state.cr_edit_description = st.text_area(
                    "Change Request Description:",
                    value=st.session_state.cr_edit_description,
                    height=250
                )
                st.divider()
                col1, col2, _ = st.columns([1, 1, 5])
                with col1:
                    if st.button("Save Changes", use_container_width=True, type="primary"):
                        if st.session_state.orchestrator.save_edited_change_request(st.session_state.cr_edit_description):
                            st.toast("✅ Change Request updated!")
                            del st.session_state.cr_edit_description # Clean up
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("Failed to save changes.")
                with col2:
                    if st.button("Cancel Edit", use_container_width=True):
                        st.session_state.orchestrator.cancel_cr_edit()
                        del st.session_state.cr_edit_description # Clean up
                        st.rerun()

        elif current_phase_name == "REPORTING_OPERATIONAL_BUG":
            st.header("Report Operational Bug")
            st.info("Use this form to report a bug found during normal operation. Provide a detailed description and assign a severity rating.")
            if 'bug_description' not in st.session_state:
                st.session_state.bug_description = ""
            if 'bug_severity' not in st.session_state:
                st.session_state.bug_severity = "Medium" # Default value
            st.session_state.bug_description = st.text_area(
                "Bug Description:",
                value=st.session_state.bug_description,
                height=250,
                help="Provide a detailed description of the bug, including steps to reproduce if possible."
            )
            st.session_state.bug_severity = st.selectbox(
                "Severity:",
                options=["Minor", "Medium", "Major"],
                index=1, # Default to 'Medium'
                help="Major: An entire function doesn't work or the system crashes.\n\nMedium: A function is partly working or produces incorrect results.\n\nMinor: A small inconvenience or cosmetic issue; nothing is broken."
            )
            st.divider()
            col1, col2, _ = st.columns([1, 1, 5])
            with col1:
                if st.button("Save Bug Report", use_container_width=True, type="primary"):
                    if st.session_state.bug_description.strip():
                        if st.session_state.orchestrator.save_bug_report(st.session_state.bug_description, st.session_state.bug_severity):
                            st.toast("✅ Bug Report saved!")
                            del st.session_state.bug_description
                            del st.session_state.bug_severity
                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("Failed to save the Bug Report.")
                    else:
                        st.warning("The bug description cannot be empty.")
            with col2:
                if st.button("Cancel", use_container_width=True):
                    if 'bug_description' in st.session_state:
                        del st.session_state.bug_description
                    if 'bug_severity' in st.session_state:
                        del st.session_state.bug_severity
                    st.session_state.orchestrator.set_phase("GENESIS")
                    st.rerun()

        elif current_phase_name == "AWAITING_LINKED_DELETE_CONFIRMATION":
            st.header("Confirm Linked Deletion")
            st.error("WARNING: This is a linked item. Deleting it will also delete its corresponding dependent item and may involve rolling back a specification change. This action cannot be undone.")
            linked_pair = st.session_state.orchestrator.task_awaiting_approval
            if linked_pair:
                primary_id = linked_pair.get('primary_cr_id')
                linked_id = linked_pair.get('linked_cr_id')
                st.markdown(f"You are attempting to delete **CR-{primary_id}**, which is linked to **CR-{linked_id}**.")
                st.divider()
                col1, col2, _ = st.columns([1.5, 1, 3])
                with col1:
                    if st.button("Yes, Delete Both Items", type="primary", use_container_width=True):
                        with st.spinner("Processing linked deletion..."):
                            st.session_state.orchestrator.handle_linked_delete_confirmation(primary_id, linked_id)
                        st.toast("Linked items have been deleted.")
                        st.rerun()
                with col2:
                    if st.button("No, Cancel", use_container_width=True):
                        st.session_state.orchestrator.task_awaiting_approval = None
                        st.session_state.orchestrator.set_phase("IMPLEMENTING_CHANGE_REQUEST")
                        st.rerun()
            else:
                st.error("Could not retrieve linked item details. Returning to the register.")
                if st.button("Back to Register"):
                    st.session_state.orchestrator.set_phase("IMPLEMENTING_CHANGE_REQUEST")
                    st.rerun()

        else: # Fallback for any other phase
            st.header(f"Current Phase: {current_phase_name}")
            st.info("The UI for this phase is under construction.")

elif page == "Documents":
    st.header("Project Documents")
    st.markdown("Select a project to view and download its core documents.")

    # --- Project Selection Logic ---
    doc_project_id = st.session_state.orchestrator.project_id
    doc_project_name = st.session_state.orchestrator.project_name

    project_history = st.session_state.orchestrator.get_project_history()

    if project_history:
        # Create a dictionary to map display names to project IDs
        history_options = {f"{row['project_name']} (ID: {row['project_id']})": row['project_id'] for row in project_history}

        # If a project is active, find its corresponding display name to set the default
        default_index = 0
        if doc_project_id:
            try:
                active_project_display_name = next(name for name, pid in history_options.items() if pid == doc_project_id)
                default_index = list(history_options.keys()).index(active_project_display_name) + 1
            except StopIteration:
                pass # Active project might not be in history yet

        selected_option = st.selectbox(
            "Select a Project:",
            options=[""] + list(history_options.keys()),
            index=default_index,
            help="Select any active or archived project to view its documents."
        )

        if selected_option:
            doc_project_id = history_options[selected_option]
            doc_project_name = selected_option.split(' (ID:')[0]
        else:
            # Clear the ID if the user selects the blank option
            doc_project_id = None

    elif doc_project_id:
        st.info(f"Displaying documents for the only active project: **{doc_project_name}**")
    else:
        st.warning("Please start a new project to generate documents.")
        st.stop()

    # --- Document Display and Download ---
    if doc_project_id:
        st.subheader(f"Documents for: {doc_project_name}")

        with st.session_state.orchestrator.db_manager as db:
            project_docs = db.get_project_by_id(doc_project_id)

        # Project Brief
        with st.expander("Project Brief", expanded=True): # Expanded by default
            brief_path_str = project_docs['project_brief_path'] if project_docs else None
            if brief_path_str:
                brief_path = Path(brief_path_str)
                if brief_path.exists():
                    try:
                        # Display content for quick view
                        if brief_path.suffix == '.docx':
                            doc = docx.Document(brief_path)
                            brief_content = "\n".join([p.text for p in doc.paragraphs])
                            st.text_area("Brief Content", brief_content, height=200, disabled=True, key=f"brief_view_{doc_project_id}")
                        else: # .md, .txt
                            brief_content = brief_path.read_text(encoding='utf-8')
                            st.text_area("Brief Content", brief_content, height=200, disabled=True, key=f"brief_view_{doc_project_id}")

                        # Add download button for the original file
                        with open(brief_path, "rb") as f:
                            st.download_button(
                                label="📄 Download Original Brief",
                                data=f,
                                file_name=brief_path.name,
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"Error reading brief file: {e}")
                else:
                    st.warning("Brief file path is recorded in the database, but the file was not found on disk.")
            else:
                st.info("No project brief was saved for this project.")

        if project_docs:
            report_generator = ReportGeneratorAgent()

            # Application Specification
            with st.expander("Application Specification", expanded=False):
                # CORRECTED: Use dictionary-style key access for sqlite3.Row
                spec_text = project_docs['final_spec_text'] if project_docs else None
                if spec_text:
                    st.text_area("Spec Content", spec_text, height=300, disabled=True, key=f"spec_{doc_project_id}")
                    spec_docx_bytes = report_generator.generate_text_document_docx(f"Application Specification - {doc_project_name}", spec_text)
                    st.download_button("📄 Print to .docx", spec_docx_bytes, f"AppSpec_{doc_project_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"download_spec_{doc_project_id}")
                else:
                    st.info("This document has not been generated for this project yet.")

            # UX/UI Specification
            with st.expander("UX/UI Specification", expanded=False):
                ux_spec_text = project_docs['ux_spec_text'] if project_docs else None
                if ux_spec_text:
                    st.text_area("UX/UI Spec Content", ux_spec_text, height=300, disabled=True, key=f"ux_spec_{doc_project_id}")
                    ux_spec_docx_bytes = report_generator.generate_text_document_docx(f"UX-UI Specification - {doc_project_name}", ux_spec_text)
                    st.download_button("📄 Print to .docx", ux_spec_docx_bytes, f"UX-Spec_{doc_project_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"download_ux_spec_{doc_project_id}")
                else:
                    st.info("This document has not been generated for this project yet.")

            # Technical Specification
            with st.expander("Technical Specification", expanded=False):
                # CORRECTED: Use dictionary-style key access
                tech_spec_text = project_docs['tech_spec_text'] if project_docs else None
                if tech_spec_text:
                    st.text_area("Tech Spec Content", tech_spec_text, height=300, disabled=True, key=f"tech_spec_{doc_project_id}")
                    tech_spec_docx_bytes = report_generator.generate_text_document_docx(f"Technical Specification - {doc_project_name}", tech_spec_text)
                    st.download_button("📄 Print to .docx", tech_spec_docx_bytes, f"TechSpec_{doc_project_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"download_tech_spec_{doc_project_id}")
                else:
                    st.info("This document has not been generated for this project yet.")

            # Coding Standard
            with st.expander("Coding Standard", expanded=False):
                coding_standard_text = project_docs['coding_standard_text'] if project_docs else None
                if coding_standard_text:
                    st.text_area("Coding Standard Content", coding_standard_text, height=300, disabled=True, key=f"cs_{doc_project_id}")
                    cs_docx_bytes = report_generator.generate_text_document_docx(f"Coding Standard - {doc_project_name}", coding_standard_text)
                    st.download_button("📄 Print to .docx", cs_docx_bytes, f"CodingStandard_{doc_project_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"download_cs_{doc_project_id}")
                else:
                    st.info("This document has not been generated for this project yet.")

            # Development Plan
            with st.expander("Development Plan", expanded=False):
                dev_plan_text = project_docs['development_plan_text'] if project_docs else None
                if dev_plan_text:
                    # Bugfix: Strip the text header before parsing as JSON
                    try:
                        # The header is separated by a distinct line of 50 dashes
                        json_content_str = dev_plan_text.split(f"\n{'-' * 50}\n\n", 1)[1]
                        st.json(json_content_str)
                    except (IndexError, json.JSONDecodeError):
                        # Fallback to display raw text if stripping fails or content is not as expected
                        st.text_area("Dev Plan Content (Raw)", dev_plan_text, height=300, disabled=True, key=f"dev_plan_{doc_project_id}")

                    dev_plan_docx_bytes = report_generator.generate_text_document_docx(f"Development Plan - {doc_project_name}", dev_plan_text, is_code=True)
                    st.download_button("📄 Print to .docx", dev_plan_docx_bytes, f"DevPlan_{doc_project_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"download_dev_plan_{doc_project_id}")
                else:
                    st.info("This document has not been generated for this project yet.")

            # Complexity & Risk Assessment
            with st.expander("Complexity & Risk Assessment", expanded=False):
                assessment_text = project_docs['complexity_assessment_text'] if project_docs else None
                if assessment_text:
                    # The assessment is stored as a JSON string with a footnote.
                    # We can display it in a clean way.
                    try:
                        # Find the split between JSON and the footnote
                        json_end_index = assessment_text.rfind('}') + 1
                        json_part_str = assessment_text[:json_end_index]
                        note_part = assessment_text[json_end_index:]

                        # Display the JSON part in a st.json element for good formatting
                        st.json(json_part_str)
                        # Display the footnote
                        st.caption(note_part.strip())

                    except Exception:
                        # Fallback for any parsing error
                        st.text_area("Assessment Content", assessment_text, height=300, disabled=True, key=f"assess_{doc_project_id}")

                    assessment_docx_bytes = report_generator.generate_text_document_docx(f"Complexity & Risk Assessment - {doc_project_name}", assessment_text, is_code=True)
                    st.download_button("📄 Print to .docx", assessment_docx_bytes, f"Assessment_{doc_project_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"download_assess_{doc_project_id}")
                else:
                    st.info("This document has not been generated for this project yet.")

            # Integration Plan
            with st.expander("Integration Plan", expanded=False):
                # CORRECTED: Use dictionary-style key access
                integration_plan_text = project_docs['integration_plan_text'] if project_docs else None
                if integration_plan_text:
                    st.json(integration_plan_text)
                    integration_plan_docx_bytes = report_generator.generate_text_document_docx(f"Integration Plan - {doc_project_name}", integration_plan_text, is_code=True)
                    st.download_button("📄 Print to .docx", integration_plan_docx_bytes, f"IntegrationPlan_{doc_project_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"download_integration_plan_{doc_project_id}")
                else:
                    st.info("This document has not been generated for this project yet.")

            # UI Test Plan
            with st.expander("UI Test Plan", expanded=False):
                ui_test_plan_text = project_docs['ui_test_plan_text']
                if ui_test_plan_text:
                    st.markdown(ui_test_plan_text)
                    ui_test_plan_docx_bytes = report_generator.generate_text_document_docx(f"UI Test Plan - {doc_project_name}", ui_test_plan_text)
                    st.download_button("📄 Print to .docx", ui_test_plan_docx_bytes, f"UITestPlan_{doc_project_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="download_ui_test_plan")
                else:
                    st.info("This document has not been generated for this project yet.")

        else:
            st.error(f"Could not retrieve document data for project ID: {doc_project_id}")

elif page == "Reports":
    st.header("Project Reports")

    # --- Project Selection Logic ---
    report_project_id = st.session_state.orchestrator.project_id
    project_history = st.session_state.orchestrator.get_project_history()

    # If no project is active, prompt the user to select one from history
    if not report_project_id and project_history:
        st.info("No project is currently active. Please select a project from the history to view its reports.")

        history_options = {f"{row['project_name']} (ID: {row['project_id']})": row['project_id'] for row in project_history}
        selected_option = st.selectbox("Select a Project:", options=[""] + list(history_options.keys()))

        if selected_option:
            report_project_id = history_options[selected_option]

    elif not report_project_id and not project_history:
        st.warning("Please start a new project or load an archived project to view reports.")
        st.stop()

    if not report_project_id:
        st.stop() # Stop if no project is active and none is selected

    st.subheader(f"Displaying Reports for Project ID: {report_project_id}")

    # --- Instantiate the report generator agent ---
    # We create it here to ensure it's available for both reports
    try:
        from agents.agent_report_generator import ReportGeneratorAgent
        report_generator = ReportGeneratorAgent()
    except ImportError:
        st.error("ReportGeneratorAgent could not be loaded. Please check the 'agents' directory.")
        st.stop()


    # --- Report 1: Development Progress Summary ---
    st.divider()
    st.markdown("#### Development Progress Summary")
    with st.session_state.orchestrator.db_manager as db:
        all_artifacts = db.get_all_artifacts_for_project(report_project_id)
        status_counts = db.get_component_counts_by_status(report_project_id)

    if not all_artifacts:
        st.info("No components have been defined for this project yet.")
    else:
        total_components = len(all_artifacts)
        st.metric(label="Total Components Defined", value=total_components)

        # Display status counts in columns for a cleaner look
        cols = st.columns(len(status_counts) or 1)
        for i, (status, count) in enumerate(status_counts.items()):
            cols[i].metric(label=status.replace("_", " ").title(), value=count)

        # Download button for this report
        summary_docx_bytes = report_generator.generate_progress_summary_docx(total_components, status_counts)
        st.download_button(
            label="📄 Print to .docx",
            data=summary_docx_bytes,
            file_name=f"ASDF_Progress_Summary_{report_project_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # --- Report 2: Change Requests & Bug Fixes ---
    st.divider()
    st.markdown("#### Change Requests & Bug Fixes")

    filter_option = st.selectbox("Filter by:", options=["Pending", "Closed", "All"])

    report_data = st.session_state.orchestrator.get_cr_and_bug_report_data(report_project_id, filter_option)

    if not report_data:
        st.info(f"No items found for the '{filter_option}' filter.")
    else:
        df = pd.DataFrame(report_data)
        st.dataframe(df, use_container_width=True, hide_index=True)

        # Download button for this report
        cr_bug_docx_bytes = report_generator.generate_cr_bug_report_docx(report_data, filter_option)
        st.download_button(
            label="📄 Print to .docx",
            data=cr_bug_docx_bytes,
            file_name=f"ASDF_CR_Bug_Report_{report_project_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="cr_bug_download"
        )

elif page == "Settings":
    st.header("Factory Settings")
    st.markdown("Configure the operational parameters of the ASDF.")
    st.divider()

    # --- LLM Service Configuration ---
    st.subheader("LLM Service Configuration")

    with st.session_state.orchestrator.db_manager as db:
        all_config = db.get_all_config_values()

    provider_options = ["Gemini", "ChatGPT", "Claude", "Phi-3 (Local)", "Any Other"]
    current_provider = all_config.get("SELECTED_LLM_PROVIDER", "Gemini")

    try:
        provider_index = provider_options.index(current_provider)
    except ValueError:
        provider_index = 0

    selected_provider = st.selectbox(
        "Select LLM Provider",
        options=provider_options,
        index=provider_index,
        key="selected_llm_provider"
    )

    if selected_provider == "Gemini":
        st.text_input("Gemini API Key", type="password", key="gemini_api_key", value=all_config.get("GEMINI_API_KEY", ""))
        st.text_input("Reasoning Model", key="gemini_reasoning_model", value=all_config.get("GEMINI_REASONING_MODEL", "gemini-2.5-pro"))
        st.text_input("Fast Model", key="gemini_fast_model", value=all_config.get("GEMINI_FAST_MODEL", "gemini-2.5-flash-preview-05-20"))
    elif selected_provider == "ChatGPT":
        st.text_input("OpenAI API Key", type="password", key="openai_api_key", value=all_config.get("OPENAI_API_KEY", ""))
        st.text_input("Reasoning Model", key="openai_reasoning_model", value=all_config.get("OPENAI_REASONING_MODEL", "gpt-4-turbo"))
        st.text_input("Fast Model", key="openai_fast_model", value=all_config.get("OPENAI_FAST_MODEL", "gpt-3.5-turbo"))
    elif selected_provider == "Claude":
        st.text_input("Anthropic API Key", type="password", key="anthropic_api_key", value=all_config.get("ANTHROPIC_API_KEY", ""))
        st.text_input("Reasoning Model", key="anthropic_reasoning_model", value=all_config.get("ANTHROPIC_REASONING_MODEL", "claude-3-opus-20240229"))
        st.text_input("Fast Model", key="anthropic_fast_model", value=all_config.get("ANTHROPIC_FAST_MODEL", "claude-3-haiku-20240307"))
    elif selected_provider == "Phi-3 (Local)":
        st.info("No configuration needed. Ensure your local Ollama server is running with the 'phi3' model available.")
    elif selected_provider == "Any Other":
        st.caption("Use this for any other publicly available model, or for your own company's private LLM.")
        st.text_input("Endpoint URL", key="custom_endpoint_url", value=all_config.get("CUSTOM_ENDPOINT_URL", ""))
        st.text_input("Endpoint API Key", type="password", key="custom_endpoint_api_key", value=all_config.get("CUSTOM_ENDPOINT_API_KEY", ""))
        st.text_input("Reasoning Model", key="custom_reasoning_model", value=all_config.get("CUSTOM_REASONING_MODEL", ""))
        st.text_input("Fast Model", key="custom_fast_model", value=all_config.get("CUSTOM_FAST_MODEL", ""))

    st.caption("Pro-Tip: For any provider, you can use the fast model name in both fields to optimize for speed and cost.")

    def save_llm_settings():
        new_provider = st.session_state.selected_llm_provider

        if st.session_state.orchestrator.project_id:
            with st.session_state.orchestrator.db_manager as db:
                all_config = db.get_all_config_values()
                current_active_limit = int(all_config.get("CONTEXT_WINDOW_CHAR_LIMIT", "0"))
                current_provider = all_config.get("SELECTED_LLM_PROVIDER")

                provider_key_map = {
                    "Gemini": "GEMINI_CONTEXT_LIMIT", "ChatGPT": "OPENAI_CONTEXT_LIMIT",
                    "Claude": "ANTHROPIC_CONTEXT_LIMIT", "Phi-3 (Local)": "LOCALPHI3_CONTEXT_LIMIT",
                    "Any Other": "ENTERPRISE_CONTEXT_LIMIT"
                }
                new_provider_default_key = provider_key_map.get(new_provider)
                all_config.update(db.get_all_config_values())
                new_provider_default_limit = int(all_config.get(new_provider_default_key, "0"))

                if new_provider != current_provider and new_provider_default_limit < current_active_limit:
                    logging.info(f"Triggering re-assessment: New limit {new_provider_default_limit} is less than current {current_active_limit}.")
                    st.session_state.reassessment_required = True
                    st.session_state.pending_llm_provider = new_provider
                    st.session_state.previous_llm_provider = current_provider
                    st.session_state.orchestrator.set_phase("AWAITING_REASSESSMENT_CONFIRMATION")
                    st.warning("Re-assessment required due to smaller context window. Please navigate to the Project page to proceed.")
                    st.rerun()
                    return

        provider = new_provider

        if provider in ["Gemini", "ChatGPT", "Claude", "Any Other"]:
            prefix_map = {"Gemini": "gemini", "ChatGPT": "openai", "Claude": "anthropic", "Any Other": "custom"}
            prefix = prefix_map.get(provider)
            reasoning_key = f"{prefix}_reasoning_model"
            fast_key = f"{prefix}_fast_model"
            reasoning_val = st.session_state.get(reasoning_key, "").strip()
            fast_val = st.session_state.get(fast_key, "").strip()
            if not reasoning_val and not fast_val:
                st.warning(f"Please enter a model name for 'Reasoning Model' and/or 'Fast Model' for {provider}.")
                return
            if not reasoning_val: st.session_state[reasoning_key] = fast_val
            elif not fast_val: st.session_state[fast_key] = reasoning_val

        settings_to_save = {"SELECTED_LLM_PROVIDER": provider}

        if provider == "Gemini":
            settings_to_save["GEMINI_API_KEY"] = st.session_state.gemini_api_key
            settings_to_save["GEMINI_REASONING_MODEL"] = st.session_state.gemini_reasoning_model
            settings_to_save["GEMINI_FAST_MODEL"] = st.session_state.gemini_fast_model
        elif provider == "ChatGPT":
            settings_to_save["OPENAI_API_KEY"] = st.session_state.openai_api_key
            settings_to_save["OPENAI_REASONING_MODEL"] = st.session_state.openai_reasoning_model
            settings_to_save["OPENAI_FAST_MODEL"] = st.session_state.openai_fast_model
        elif provider == "Claude":
            settings_to_save["ANTHROPIC_API_KEY"] = st.session_state.anthropic_api_key
            settings_to_save["ANTHROPIC_REASONING_MODEL"] = st.session_state.anthropic_reasoning_model
            settings_to_save["ANTHROPIC_FAST_MODEL"] = st.session_state.anthropic_fast_model
        elif provider == "Any Other":
            settings_to_save["CUSTOM_ENDPOINT_URL"] = st.session_state.custom_endpoint_url
            settings_to_save["CUSTOM_ENDPOINT_API_KEY"] = st.session_state.custom_endpoint_api_key
            settings_to_save["CUSTOM_REASONING_MODEL"] = st.session_state.custom_reasoning_model
            settings_to_save["CUSTOM_FAST_MODEL"] = st.session_state.custom_fast_model

        try:
            with st.session_state.orchestrator.db_manager as db:
                for key, value in settings_to_save.items():
                    db.set_config_value(key, str(value))

                provider_key_map = {
                    "Gemini": "GEMINI_CONTEXT_LIMIT", "ChatGPT": "OPENAI_CONTEXT_LIMIT",
                    "Claude": "ANTHROPIC_CONTEXT_LIMIT", "Phi-3 (Local)": "LOCALPHI3_CONTEXT_LIMIT",
                    "Any Other": "ENTERPRISE_CONTEXT_LIMIT"
                }
                provider_default_key = provider_key_map.get(provider)
                if provider_default_key:
                    provider_default_value = db.get_config_value(provider_default_key)
                    if provider_default_value:
                        db.set_config_value("CONTEXT_WINDOW_CHAR_LIMIT", provider_default_value)
                        logging.info(f"Active context limit updated to default for {provider}: {provider_default_value} chars.")

            st.session_state.orchestrator.llm_service = st.session_state.orchestrator._create_llm_service()
            if st.session_state.orchestrator.llm_service:
                st.success("✅ LLM Service settings saved and activated!")
            else:
                st.error("Settings saved, but failed to activate the LLM Service. Please check your API key or endpoint details.")
        except Exception as e:
            st.error(f"Failed to save LLM settings: {e}")

    st.button("Save LLM Service Settings", on_click=save_llm_settings, use_container_width=True, type="primary")

    st.divider()

    st.subheader("Additional Settings")

    with st.session_state.orchestrator.db_manager as db:
        all_config = db.get_all_config_values()

    st.number_input(
        "Maximum Automated Debug Attempts",
        min_value=1, key="max_debug_attempts",
        value=int(all_config.get("MAX_DEBUG_ATTEMPTS", 2)),
        help="Defines the number of automated fix attempts the Debug Pipeline will perform before escalating to the PM."
    )

    # Get the provider currently selected in the dropdown (might not be saved yet)
    selected_provider_in_ui = st.session_state.get("selected_llm_provider", current_provider)
    # Get the provider that is actually saved in the database
    saved_provider_in_db = all_config.get("SELECTED_LLM_PROVIDER", "Gemini")

    display_value = 0
    if selected_provider_in_ui == saved_provider_in_db:
        # If the selection matches what's saved, display the saved active limit.
        display_value = int(all_config.get("CONTEXT_WINDOW_CHAR_LIMIT", 2000000))
    else:
        # If the user has selected a *new* provider, predictively show its default.
        provider_key_map = {
            "Gemini": "GEMINI_CONTEXT_LIMIT", "ChatGPT": "OPENAI_CONTEXT_LIMIT",
            "Claude": "ANTHROPIC_CONTEXT_LIMIT", "Phi-3 (Local)": "LOCALPHI3_CONTEXT_LIMIT",
            "Any Other": "ENTERPRISE_CONTEXT_LIMIT"
        }
        provider_default_key = provider_key_map.get(selected_provider_in_ui)
        # Get the default value from the full config dictionary loaded at the top of the page.
        display_value = int(all_config.get(provider_default_key, 2000000))

    st.number_input(
        "Context Window Character Limit",
        min_value=10000,
        step=5000,
        key="context_window_limit",
        value=display_value, # Use the dynamically determined value
        help="Defines the maximum number of characters to send to the LLM for complex analysis. This value defaults to the recommended limit for the selected provider."
    )

    pm_checkpoint_options = {"ALWAYS_ASK": "Always ask before proceeding", "AUTO_PROCEED": "Automatically proceed if successful"}
    current_pm_behavior = all_config.get("PM_CHECKPOINT_BEHAVIOR", "ALWAYS_ASK")
    pm_checkpoint_index = list(pm_checkpoint_options.keys()).index(current_pm_behavior)
    st.selectbox(
        "PM Checkpoint Behavior (Genesis Phase)",
        options=pm_checkpoint_options.values(),
        index=pm_checkpoint_index, key="pm_checkpoint_behavior",
        help="Controls the factory's behavior after successfully developing a component."
    )

    logging_options = ["Standard", "Detailed", "Debug"]
    current_logging_level = all_config.get("LOGGING_LEVEL", "Standard")
    try:
        logging_index = logging_options.index(current_logging_level)
    except ValueError:
        logging_index = 0
    st.selectbox(
        "ASDF Operational Logging Level",
        options=logging_options, index=logging_index, key="logging_level",
        help="Controls the verbosity of ASDF's internal logs, useful for troubleshooting the factory application itself."
    )

    st.text_input("Default Base Path for New Target Projects", key="default_project_path", value=all_config.get("DEFAULT_PROJECT_PATH", ""), help="Optional. Set a default parent directory (e.g., 'C:\\Users\\YourName\\Projects').")
    st.text_input("Default Project Archive Path", key="default_archive_path", value=all_config.get("DEFAULT_ARCHIVE_PATH", ""), help="Optional. Set a default folder for saving project archives.")

    def save_additional_settings():
        settings_to_save = {
            "MAX_DEBUG_ATTEMPTS": st.session_state.max_debug_attempts,
            "CONTEXT_WINDOW_CHAR_LIMIT": st.session_state.context_window_limit,
            "LOGGING_LEVEL": st.session_state.logging_level,
            "DEFAULT_PROJECT_PATH": st.session_state.default_project_path,
            "DEFAULT_ARCHIVE_PATH": st.session_state.default_archive_path
        }
        selected_pm_behavior_value = st.session_state.pm_checkpoint_behavior
        for key, value in pm_checkpoint_options.items():
            if value == selected_pm_behavior_value:
                settings_to_save["PM_CHECKPOINT_BEHAVIOR"] = key
                break
        with st.session_state.orchestrator.db_manager as db:
            for key, value in settings_to_save.items():
                db.set_config_value(key, str(value))
        st.success("✅ Additional settings saved!")

    st.button("Save All Additional Settings", on_click=save_additional_settings, type="primary", use_container_width=True)
